import tempfile
import unittest
from contextlib import closing
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document

import ihale_belge_arsivi
import veritabani


class IhaleBelgeArsiviTesti(unittest.TestCase):
    def test_meslek_lisesi_ayri_okul_turu_olarak_belirlenir(self):
        sonuc = ihale_belge_arsivi.metin_verilerini_cikar("""
            Atatürk Mesleki ve Teknik Anadolu Lisesi kantin ihalesi
            Öğrenci sayısı: 900
            Aylık muhammen kira bedeli: 30.000 TL
        """)
        self.assertEqual("Meslek Lisesi", sonuc["okul_turu"])

    def test_ocr_kelime_sirasi_ve_ogrenci_personel_toplami_okunur(self):
        sonuc = ihale_belge_arsivi.metin_verilerini_cikar("""
            Şehit Oğuz Özgür Çevik Anadolu Lisesi
            Öğrenci sayısı: 705+48 (öğrenci+personel)
            İhaleye esas aylıkmuhammenkirabedeli: 90.000,00 TL
            Yıllık muhammen kira bedeli: 810.000,00 TL (9 ay)
        """)
        self.assertEqual(705, sonuc["ogrenci_sayisi"])
        self.assertEqual(48, sonuc["personel_sayisi"])
        self.assertEqual(90000, sonuc["muhammen_bedel_aylik"])
        self.assertEqual(810000, sonuc["muhammen_bedel_yillik"])

    def test_muhammen_bedel_aylik_sonra_yazildiginda_okunur(self):
        sonuc = ihale_belge_arsivi.metin_verilerini_cikar("""
            Salih İlhan İmam Hatip Ortaokulu
            Muhammen bedel: aylık 8.000,00 TL yıllık 72.000,00 TL
            Öğrenci mevcudu: 350
        """)
        self.assertEqual(8000, sonuc["muhammen_bedel_aylik"])
        self.assertEqual(72000, sonuc["muhammen_bedel_yillik"])

    def test_ucret_ve_okul_verilerini_ayri_ayri_cikarir(self):
        sonuc = ihale_belge_arsivi.metin_verilerini_cikar("""
            Atatürk İlkokulu kantin ihalesi
            Öğrenci sayısı: 742
            Öğretmen ve diğer personel sayısı: 61
            Kantin alanı: 48,5 m²
            Aylık muhammen kira bedeli: 18.750,00 TL
            Geçici teminat: en az 6.750,00 TL
            İhale şartnamesi bedeli: 2.500 TL
            Kira süresi: 1 yıl
        """)
        self.assertEqual(742, sonuc["ogrenci_sayisi"])
        self.assertEqual(61, sonuc["personel_sayisi"])
        self.assertEqual(18750, sonuc["muhammen_bedel"])
        self.assertEqual(18750, sonuc["muhammen_bedel_aylik"])
        self.assertEqual(168750, sonuc["muhammen_bedel_yillik"])
        self.assertEqual("aylik", sonuc["muhammen_bedel_donemi"])
        self.assertEqual("İlkokul", sonuc["okul_turu"])
        self.assertEqual(6750, sonuc["gecici_teminat"])
        self.assertEqual(2500, sonuc["sartname_bedeli"])
        self.assertEqual(48.5, sonuc["kantin_alani_m2"])
        self.assertEqual(12, sonuc["kira_suresi_ay"])

    def test_parantezli_muhammen_ve_tutar_once_yazilan_teminat_okunur(self):
        sonuc = ihale_belge_arsivi.metin_verilerini_cikar("""
            MUHAMMEN BEDEL 8.000 (sekiz bin)- TL
            Teklifin yüzde üçünden az olmamak üzere (en az 2.040,00 TL)
            isteklice geçici teminat yatırılır.
            İhale şartnamesi bedeli 2.500 TL olarak yatırılır.
        """)
        self.assertEqual(8000, sonuc["muhammen_bedel"])
        self.assertEqual(2040, sonuc["gecici_teminat"])
        self.assertEqual(2500, sonuc["sartname_bedeli"])

    def test_docx_yerele_kaydedilir_ve_ana_ilana_baglanir(self):
        with tempfile.TemporaryDirectory() as klasor:
            db = Path(klasor) / "arsiv.db"
            arsiv = Path(klasor) / "belgeler"
            with patch.object(veritabani, "DB", db), patch.object(
                ihale_belge_arsivi, "ARSIV_KOK", arsiv
            ):
                veritabani.tablo_olustur()
                with closing(veritabani.baglan()) as conn, conn:
                    kaynak_id = conn.execute("""
                        INSERT INTO kaynaklar(
                            kurum_adi, il, ilce, url, eklenme_tarihi
                        ) VALUES ('Mamak MEM', 'Ankara', 'Mamak',
                                  'https://mamak.meb.gov.tr', '2026-07-01')
                    """).lastrowid
                    aday_id = conn.execute("""
                        INSERT INTO duyuru_adaylari(
                            kaynak_id, baslik, url, dosya_turu, eslesme_turu,
                            yayin_tarihi, ilk_gorulme, son_gorulme
                        ) VALUES (?, 'Atatürk İlkokulu Kantin İhalesi',
                                  'https://mamak.meb.gov.tr/ilan.docx', '.docx',
                                  'dosya', '2026-07-01', '2026-07-01', '2026-07-01')
                    """, (kaynak_id,)).lastrowid

                belge = Document()
                belge.add_paragraph("Öğrenci sayısı: 500")
                belge.add_paragraph("Aylık muhammen bedel: 9.000 TL")
                belge.add_paragraph("Şartname bedeli: 750 TL")
                akim = BytesIO()
                belge.save(akim)
                sonuc = ihale_belge_arsivi.icerigi_arsivle(
                    aday_id,
                    "https://mamak.meb.gov.tr/ilan.docx",
                    akim.getvalue(),
                    yayin_tarihi="2026-07-01",
                    baslik="Atatürk İlkokulu Kantin İhalesi",
                )

                self.assertEqual("analiz_edildi", sonuc["durum"])
                self.assertTrue(Path(sonuc["yerel_yol"]).is_file())
                with closing(veritabani.baglan()) as conn:
                    analiz = conn.execute("""
                        SELECT okul_adi, okul_turu, ogrenci_sayisi,
                               muhammen_bedel_aylik, muhammen_bedel_yillik,
                               sartname_bedeli
                        FROM ilan_analiz_verileri WHERE aday_id=?
                    """, (aday_id,)).fetchone()
                self.assertEqual(500, analiz["ogrenci_sayisi"])
                self.assertEqual("İlkokul", analiz["okul_turu"])
                self.assertEqual(9000, analiz["muhammen_bedel_aylik"])
                self.assertEqual(81000, analiz["muhammen_bedel_yillik"])
                self.assertEqual(750, analiz["sartname_bedeli"])

    def test_yillik_muhammen_aylik_analiz_degerine_donusturulur(self):
        sonuc = ihale_belge_arsivi.metin_verilerini_cikar("""
            Gazi Anadolu Lisesi kantin ihalesi
            Öğrenci sayısı: 800
            Yıllık muhammen kira bedeli: 240.000,00 TL
        """)
        self.assertEqual("Lise", sonuc["okul_turu"])
        self.assertEqual(240000, sonuc["muhammen_bedel_yillik"])
        self.assertAlmostEqual(26666.67, sonuc["muhammen_bedel_aylik"], places=2)
        self.assertEqual("yillik", sonuc["muhammen_bedel_donemi"])

    def test_toplu_ekler_okul_kaydina_tek_ek_ana_ilana_baglanir(self):
        with tempfile.TemporaryDirectory() as klasor, patch.object(
            veritabani, "DB", Path(klasor) / "kuyruk.db"
        ):
            veritabani.tablo_olustur()
            with closing(veritabani.baglan()) as conn, conn:
                kaynak_id = conn.execute("""
                    INSERT INTO kaynaklar(kurum_adi, il, ilce, url, eklenme_tarihi)
                    VALUES ('Test MEM', 'Ankara', 'Mamak',
                            'https://test.meb.gov.tr', '2026-07-01')
                """).lastrowid
                parent = conn.execute("""
                    INSERT INTO duyuru_adaylari(
                        kaynak_id, baslik, url, eslesme_turu, yayin_tarihi,
                        ilk_gorulme, son_gorulme
                    ) VALUES (?, 'İki okul kantin ihalesi',
                              'https://test.meb.gov.tr/toplu', 'detay',
                              '2026-07-01', '2026-07-01', '2026-07-01')
                """, (kaynak_id,)).lastrowid
                cocuklar = []
                for sira in (1, 2):
                    cocuklar.append(conn.execute("""
                        INSERT INTO duyuru_adaylari(
                            kaynak_id, baslik, url, detay_url, eslesme_turu,
                            dosya_turu, yayin_tarihi, ilk_gorulme, son_gorulme
                        ) VALUES (?, ?, ?, 'https://test.meb.gov.tr/toplu',
                                  'ek_dosya', '.pdf', '2026-07-01',
                                  '2026-07-01', '2026-07-01')
                    """, (
                        kaynak_id, f'{sira}. okul',
                        f'https://test.meb.gov.tr/{sira}.pdf',
                    )).lastrowid)
                tek_parent = conn.execute("""
                    INSERT INTO duyuru_adaylari(
                        kaynak_id, baslik, url, eslesme_turu, yayin_tarihi,
                        ilk_gorulme, son_gorulme
                    ) VALUES (?, 'Tek okul', 'https://test.meb.gov.tr/tek',
                              'detay', '2026-07-01', '2026-07-01', '2026-07-01')
                """, (kaynak_id,)).lastrowid
                conn.execute("""
                    INSERT INTO duyuru_adaylari(
                        kaynak_id, baslik, url, detay_url, eslesme_turu,
                        dosya_turu, yayin_tarihi, ilk_gorulme, son_gorulme
                    ) VALUES (?, 'Tek okul belgesi',
                              'https://test.meb.gov.tr/tek.pdf',
                              'https://test.meb.gov.tr/tek', 'ek_dosya', '.pdf',
                              '2026-07-01', '2026-07-01', '2026-07-01')
                """, (kaynak_id,))

            kuyruk = {
                kayit["url"]: kayit["aday_id"]
                for kayit in ihale_belge_arsivi.arsiv_kuyrugu(limit=10)
            }
            self.assertEqual(cocuklar[0], kuyruk['https://test.meb.gov.tr/1.pdf'])
            self.assertEqual(cocuklar[1], kuyruk['https://test.meb.gov.tr/2.pdf'])
            self.assertEqual(tek_parent, kuyruk['https://test.meb.gov.tr/tek.pdf'])


if __name__ == "__main__":
    unittest.main()
