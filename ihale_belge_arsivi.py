"""Resmî MEB ihale belgelerini yerelde saklar ve kart verilerini çıkarır."""

from __future__ import annotations

from contextlib import closing
from datetime import date, datetime
from hashlib import sha256
from io import BytesIO
from pathlib import Path, PurePosixPath
from functools import lru_cache
from threading import Lock
from urllib.parse import unquote, urlparse
import json
import mimetypes
import re
import unicodedata

import requests
import cv2
import numpy as np
import pypdfium2 as pdfium
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from rapidocr_onnxruntime import RapidOCR

from istatistik_motoru import okul_adi_ayikla
from okul_adi_servisi import okul_adi_temizle
from veritabani import baglan, ihale_tarih_siniri


KOK = Path(__file__).resolve().parent
ARSIV_KOK = KOK / "ilanlar" / "resmi_arsiv"
DURUM_DOSYASI = KOK / "belge_arsiv_durumu.json"
AZAMI_DOSYA_BOYUTU = 20_000_000
DESTEKLENEN_UZANTILAR = {
    ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".jpg", ".jpeg", ".png", ".webp"
}
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 Chrome/126.0 Safari/537.36 Ihalex/1.0"
    )
}
_OCR_KILIDI = Lock()


class BelgeArsivHatasi(RuntimeError):
    """Belge güvenlik veya erişim kontrolünden geçemediğinde oluşur."""


def resmi_meb_url(url: object) -> bool:
    host = (urlparse(str(url or "")).hostname or "").lower()
    return host == "meb.gov.tr" or host.endswith(".meb.gov.tr")


def _guvenli_dosya_adi(url: str, ozet: str) -> str:
    ham = unquote(PurePosixPath(urlparse(url).path).name) or "ihale_belgesi"
    uzanti = Path(ham).suffix.lower()
    govde = Path(ham).stem
    govde = re.sub(r"[^0-9A-Za-zÇĞİÖŞÜçğıöşü._-]+", "_", govde).strip("._-")
    return f"{ozet[:16]}_{(govde or 'ihale_belgesi')[:110]}{uzanti}"


def _para_cevir(deger: object) -> float | None:
    metin = re.sub(r"[^0-9.,]", "", str(deger or ""))
    if not metin:
        return None
    if "," in metin:
        metin = metin.replace(".", "").replace(",", ".")
    elif metin.count(".") > 1 or (
        "." in metin and len(metin.rsplit(".", 1)[1]) == 3
    ):
        metin = metin.replace(".", "")
    try:
        sonuc = float(metin)
    except ValueError:
        return None
    return sonuc if 0 <= sonuc <= 1_000_000_000 else None


def _ascii_tr(deger: object) -> str:
    metin = unicodedata.normalize("NFKD", str(deger or "").casefold()).replace("ı", "i")
    return "".join(harf for harf in metin if not unicodedata.combining(harf))


def _sayi_bul(metin: str, desenler: tuple[str, ...], alt: int, ust: int) -> int | None:
    for desen in desenler:
        eslesme = re.search(desen, metin, re.IGNORECASE | re.DOTALL)
        if eslesme is None:
            eslesme = re.search(
                _ascii_tr(desen), _ascii_tr(metin), re.IGNORECASE | re.DOTALL
            )
        if eslesme:
            sayi = int(eslesme.group(1))
            if alt <= sayi <= ust:
                return sayi
    return None


def _para_bul(metin: str, desenler: tuple[str, ...]) -> float | None:
    for desen in desenler:
        eslesme = re.search(desen, metin, re.IGNORECASE | re.DOTALL)
        if eslesme is None:
            eslesme = re.search(
                _ascii_tr(desen), _ascii_tr(metin), re.IGNORECASE | re.DOTALL
            )
        if eslesme:
            tutar = _para_cevir(eslesme.group(1))
            if tutar is not None:
                return tutar
    return None


def okul_turu_belirle(*degerler: object) -> str | None:
    """Okul adı, belge metni veya başlıktan doğrulanabilir okul türünü döndürür."""
    metin = unicodedata.normalize(
        "NFKD", " ".join(str(deger or "") for deger in degerler).casefold()
    ).replace("ı", "i")
    metin = "".join(harf for harf in metin if not unicodedata.combining(harf))
    turler: list[str] = []
    if "ilkokul" in metin:
        turler.append("İlkokul")
    if "ortaokul" in metin:
        turler.append("Ortaokul")
    meslek_lisesi = (
        "meslek lises" in metin or "mesleki ve teknik" in metin
        or "mesleki teknik" in metin
    )
    if meslek_lisesi:
        turler.append("Meslek Lisesi")
    elif re.search(r"\blise(?:si|leri)?\b", metin):
        turler.append("Lise")
    benzersiz = list(dict.fromkeys(turler))
    if len(benzersiz) == 1:
        return benzersiz[0]
    if len(benzersiz) > 1:
        return "Karma"
    return None


def metin_verilerini_cikar(metin: str, baslik: str = "") -> dict[str, object]:
    """Belge metninden birbirine karıştırılmayan ihale alanlarını çıkarır."""
    metin = str(metin or "")[:2_000_000]
    tek_satir = " ".join(metin.split())

    ogrenci = _sayi_bul(
        tek_satir,
        (
            r"öğrenci\s+sayısı(?:\s+ve\s+öğrenim\s+durumu)?\s*[|:\-]?\s*(\d{2,5})",
            r"öğrenci\s+mevcudu\s*[|:\-]?\s*(\d{2,5})",
            r"(\d{2,5})\s+öğrenci\b",
        ),
        20,
        30_000,
    )
    personel = _sayi_bul(
        tek_satir,
        (
            r"öğrenci\s+sayısı\s*[|:\-]?\s*\d{1,5}\s*\+\s*(\d{1,4})"
            r".{0,60}?öğrenci\s*\+\s*personel",
            r"öğretmen\s+ve\s+diğer\s+personel\s+sayısı\s*[|:\-]?\s*(\d{1,4})",
            r"personel\s+sayısı\s*[|:\-]?\s*(\d{1,4})",
        ),
        1,
        5_000,
    )
    aylik_muhammen = _para_bul(
        tek_satir,
        (
            r"muhammen\s+(?:kira\s+)?bedeli?\s*[|:\-]?\s*(?:aylık|bir\s+aylık)"
            r"\s*([\d.,]+)(?:\s*\([^)]*\))?\s*[-–]?\s*(?:tl|₺)",
            r"(?:aylık|bir\s+aylık)\s*muhammen\s*(?:kira\s*)?bedeli?\s*[|:\-]?"
            r"\s*([\d.,]+)(?:\s*\([^)]*\))?\s*[-–]?\s*(?:tl|₺)",
            r"muhammen\s+(?:kira\s+)?bedeli?\s*[|:\-]?\s*([\d.,]+)"
            r".{0,140}?yıllık\s*[\d.,]+\s*ay\s*[x×*]",
            r"(?:aylık|bir\s+aylık)\s+muhammen\s+(?:kira\s+)?bedeli?\s*[|:\-]?\s*([\d.,]+)"
            r"(?:\s*\([^)]*\))?\s*[-–]?\s*(?:tl|₺)",
            r"muhammen\s+(?:kira\s+)?bedeli?\s*\(?\s*aylık\s*\)?\s*[|:\-]?\s*([\d.,]+)"
            r"(?:\s*\([^)]*\))?\s*[-–]?\s*(?:tl|₺)",
        ),
    )
    yillik_muhammen = _para_bul(
        tek_satir,
        (
            r"(?:yıllık|bir\s+yıllık|1\s+yıllık)\s+muhammen\s+(?:kira\s+)?bedeli?\s*[|:\-]?\s*([\d.,]+)"
            r"(?:\s*\([^)]*\))?\s*[-–]?\s*(?:tl|₺)",
            r"muhammen\s+(?:kira\s+)?bedeli?\s*\(?\s*yıllık\s*\)?\s*[|:\-]?\s*([\d.,]+)"
            r"(?:\s*\([^)]*\))?\s*[-–]?\s*(?:tl|₺)",
        ),
    )
    muhammen = _para_bul(
        tek_satir,
        (
            r"(?:aylık|yıllık|bir\s+(?:aylık|yıllık)|1\s+yıllık)?\s*"
            r"muhammen\s+(?:kira\s+)?bedeli?\s*[|:\-]?\s*([\d.,]+)"
            r"(?:\s*\([^)]*\))?\s*[-–]?\s*(?:tl|₺)",
        ),
    )
    muhammen_donemi = None
    if aylik_muhammen is not None:
        muhammen = aylik_muhammen
        muhammen_donemi = "aylik"
        if yillik_muhammen is None:
            # Okul kantini kira dönemi fiilî eğitim dönemi olarak 9 ay kabul edilir.
            yillik_muhammen = round(aylik_muhammen * 9, 2)
    elif yillik_muhammen is not None:
        muhammen = yillik_muhammen
        muhammen_donemi = "yillik"
        aylik_muhammen = round(yillik_muhammen / 9, 2)
    sartname = _para_bul(
        tek_satir,
        (
            r"(?:ihale\s+)?şartname(?:si)?\s+bedeli\s*[|:\-]?\s*([\d.,]+)\s*(?:tl|₺)",
            r"şartname\s*[|:\-]?\s*([\d.,]+)\s*(?:tl|₺)",
        ),
    )
    teminat = _para_bul(
        tek_satir,
        (
            r"en\s+az\s+([\d.]+,\d{2})\s*(?:tl|₺).{0,220}?geçici\s+teminat",
            r"geçici\s+teminat(?:(?!şartname|muhammen|kesin\s+teminat).){0,220}?"
            r"(?:en\s+az\s*)?([\d.]+,\d{2})\s*(?:tl|₺)",
            r"geçici\s+teminat(?:(?!şartname|muhammen|kesin\s+teminat).){0,160}?"
            r"([\d.]+)\s*(?:tl|₺)",
        ),
    )
    alan = _para_bul(
        tek_satir,
        (
            r"kantin(?:in)?\s+(?:yeri\s+ve\s+)?(?:alanı|m[²2]'?si)\s*[|:\-]?\s*([\d.,]+)\s*(?:m2|m²)",
            r"([\d.,]+)\s*(?:m2|m²)\s+(?:kantin|büfe)",
        ),
    )
    if alan is not None and not 2 <= alan <= 5_000:
        alan = None

    okul = okul_adi_ayikla(baslik)
    okul_eslesmesi = re.search(
        r"([A-ZÇĞİÖŞÜ0-9][A-ZÇĞİÖŞÜa-zçğıöşü0-9 .,'’\-]{2,110}?"
        r"(?:İlkokulu|Ortaokulu|Anadolu Lisesi|İmam Hatip Lisesi|Lisesi))",
        tek_satir,
        re.IGNORECASE,
    )
    if okul is None and okul_eslesmesi:
        okul = " ".join(okul_eslesmesi.group(1).split()).strip(" |:-")
    okul = okul_adi_temizle(okul)
    okul_turu = (
        okul_turu_belirle(okul)
        or okul_turu_belirle(baslik)
        or okul_turu_belirle(tek_satir)
    )

    adres = None
    for desen in (
        r"(?:okulun|kantinin|idarenin)?\s*adresi\s*[|:\-]\s*([^\n]{10,220})",
        r"adres\s*[|:\-]\s*([^\n]{10,220})",
    ):
        eslesme = re.search(desen, metin, re.IGNORECASE)
        if eslesme:
            adres = " ".join(eslesme.group(1).split()).strip(" |:-")[:220]
            break

    kira_suresi_ay = None
    sure = re.search(r"(?:kira|sözleşme)\s+süresi\s*[|:\-]?\s*(\d{1,2})\s*(yıl|ay)", tek_satir, re.IGNORECASE)
    if sure:
        kira_suresi_ay = int(sure.group(1)) * (12 if sure.group(2).casefold() == "yıl" else 1)
        if not 1 <= kira_suresi_ay <= 120:
            kira_suresi_ay = None

    alanlar = {
        "okul_adi": okul,
        "okul_turu": okul_turu,
        "adres": adres,
        "ogrenci_sayisi": ogrenci,
        "personel_sayisi": personel,
        "muhammen_bedel": muhammen,
        "muhammen_bedel_aylik": aylik_muhammen,
        "muhammen_bedel_yillik": yillik_muhammen,
        "muhammen_bedel_donemi": muhammen_donemi,
        "sartname_bedeli": sartname,
        "gecici_teminat": teminat,
        "kantin_alani_m2": alan,
        "kira_suresi_ay": kira_suresi_ay,
    }
    kritik = sum(alanlar[ad] is not None for ad in (
        "ogrenci_sayisi", "personel_sayisi", "muhammen_bedel_aylik",
        "sartname_bedeli", "gecici_teminat", "kantin_alani_m2",
    ))
    alanlar["belge_guveni"] = min(
        100, 30 + kritik * 10 + (5 if okul else 0) + (5 if okul_turu else 0)
    )
    alanlar["ham_metin"] = metin
    alanlar["baslik"] = baslik
    return alanlar


def _metin_kaliteli_mi(metin: str) -> bool:
    harf_sayisi = sum(harf.isalpha() for harf in str(metin or ""))
    arama = str(metin or "").casefold()
    return harf_sayisi >= 100 and any(
        anahtar in arama for anahtar in ("ihale", "kantin", "şartname", "sartname")
    )


@lru_cache(maxsize=1)
def _ocr_motoru() -> RapidOCR:
    return RapidOCR()


def _pdf_ocr_metni(icerik: bytes, sayfa_siniri: int = 6) -> str:
    pdf = pdfium.PdfDocument(BytesIO(icerik))
    parcalar: list[str] = []
    for sayfa_no in range(min(len(pdf), sayfa_siniri)):
        goruntu = pdf[sayfa_no].render(scale=1.8).to_numpy()
        with _OCR_KILIDI:
            sonuc, _ = _ocr_motoru()(goruntu)
        if sonuc:
            parcalar.append("\n".join(str(satir[1]) for satir in sonuc))
            gecici_veri = metin_verilerini_cikar("\n".join(parcalar))
            kritik_sayi = sum(
                gecici_veri.get(alan) is not None for alan in (
                    "ogrenci_sayisi", "muhammen_bedel_aylik", "sartname_bedeli",
                    "gecici_teminat", "kantin_alani_m2",
                )
            )
            if (
                gecici_veri.get("okul_adi")
                and gecici_veri.get("okul_turu")
                and gecici_veri.get("ogrenci_sayisi") is not None
                and gecici_veri.get("muhammen_bedel_aylik") is not None
            ):
                break
    return "\n".join(parcalar)


def _gorsel_ocr_metni(icerik: bytes) -> str:
    goruntu = cv2.imdecode(np.frombuffer(icerik, dtype=np.uint8), cv2.IMREAD_COLOR)
    if goruntu is None:
        return ""
    with _OCR_KILIDI:
        sonuc, _ = _ocr_motoru()(goruntu)
    return "\n".join(str(satir[1]) for satir in sonuc) if sonuc else ""


def belge_metni_oku(icerik: bytes, uzanti: str) -> str:
    uzanti = uzanti.lower()
    if uzanti == ".pdf":
        okuyucu = PdfReader(BytesIO(icerik), strict=False)
        metin = "\n".join((sayfa.extract_text() or "") for sayfa in okuyucu.pages[:40])
        return metin if _metin_kaliteli_mi(metin) else _pdf_ocr_metni(icerik)
    if uzanti == ".docx":
        belge = Document(BytesIO(icerik))
        parcalar = [p.text for p in belge.paragraphs if p.text]
        for tablo in belge.tables:
            for satir in tablo.rows:
                parcalar.append(" | ".join(hucre.text for hucre in satir.cells if hucre.text))
        return "\n".join(parcalar)
    if uzanti == ".xlsx":
        kitap = load_workbook(BytesIO(icerik), read_only=True, data_only=True)
        parcalar: list[str] = []
        for sayfa in kitap.worksheets:
            for satir in sayfa.iter_rows(values_only=True):
                parcalar.append(" | ".join(str(x) for x in satir if x is not None))
        return "\n".join(parcalar)
    if uzanti in {".jpg", ".jpeg", ".png", ".webp"}:
        return _gorsel_ocr_metni(icerik)
    return ""


def _analiz_verisini_yaz(
    conn: object,
    aday_id: int,
    belge_id: int,
    veri: dict[str, object],
) -> None:
    simdi = datetime.now().isoformat(timespec="seconds")
    conn.execute("""
        INSERT INTO ilan_analiz_verileri (
            aday_id, okul_adi, okul_turu, adres, ogrenci_sayisi, personel_sayisi,
            muhammen_bedel, muhammen_bedel_aylik, muhammen_bedel_yillik,
            muhammen_bedel_donemi, sartname_bedeli, gecici_teminat,
            kantin_alani_m2, kira_suresi_ay, belge_guveni,
            kaynak_belge_id, veri_yontemi, ham_metin, guncelleme_tarihi
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'belge_metni', ?, ?)
        ON CONFLICT(aday_id) DO UPDATE SET
            okul_adi=COALESCE(excluded.okul_adi, ilan_analiz_verileri.okul_adi),
            okul_turu=COALESCE(excluded.okul_turu, ilan_analiz_verileri.okul_turu),
            adres=COALESCE(excluded.adres, ilan_analiz_verileri.adres),
            ogrenci_sayisi=COALESCE(excluded.ogrenci_sayisi, ilan_analiz_verileri.ogrenci_sayisi),
            personel_sayisi=COALESCE(excluded.personel_sayisi, ilan_analiz_verileri.personel_sayisi),
            muhammen_bedel=COALESCE(excluded.muhammen_bedel, ilan_analiz_verileri.muhammen_bedel),
            muhammen_bedel_aylik=COALESCE(excluded.muhammen_bedel_aylik, ilan_analiz_verileri.muhammen_bedel_aylik),
            muhammen_bedel_yillik=COALESCE(excluded.muhammen_bedel_yillik, ilan_analiz_verileri.muhammen_bedel_yillik),
            muhammen_bedel_donemi=COALESCE(excluded.muhammen_bedel_donemi, ilan_analiz_verileri.muhammen_bedel_donemi),
            sartname_bedeli=COALESCE(excluded.sartname_bedeli, ilan_analiz_verileri.sartname_bedeli),
            gecici_teminat=COALESCE(excluded.gecici_teminat, ilan_analiz_verileri.gecici_teminat),
            kantin_alani_m2=COALESCE(excluded.kantin_alani_m2, ilan_analiz_verileri.kantin_alani_m2),
            kira_suresi_ay=COALESCE(excluded.kira_suresi_ay, ilan_analiz_verileri.kira_suresi_ay),
            belge_guveni=MAX(ilan_analiz_verileri.belge_guveni, excluded.belge_guveni),
            kaynak_belge_id=CASE
                WHEN excluded.belge_guveni >= ilan_analiz_verileri.belge_guveni
                THEN excluded.kaynak_belge_id ELSE ilan_analiz_verileri.kaynak_belge_id END,
            ham_metin=CASE
                WHEN LENGTH(COALESCE(excluded.ham_metin, '')) > LENGTH(COALESCE(ilan_analiz_verileri.ham_metin, ''))
                THEN excluded.ham_metin ELSE ilan_analiz_verileri.ham_metin END,
            guncelleme_tarihi=excluded.guncelleme_tarihi
    """, (
        aday_id, veri.get("okul_adi"), veri.get("okul_turu"), veri.get("adres"),
        veri.get("ogrenci_sayisi"), veri.get("personel_sayisi"),
        veri.get("muhammen_bedel"), veri.get("muhammen_bedel_aylik"),
        veri.get("muhammen_bedel_yillik"), veri.get("muhammen_bedel_donemi"),
        veri.get("sartname_bedeli"),
        veri.get("gecici_teminat"), veri.get("kantin_alani_m2"),
        veri.get("kira_suresi_ay"), int(veri.get("belge_guveni") or 0),
        belge_id, str(veri.get("ham_metin") or "")[:2_000_000], simdi,
    ))


def icerigi_arsivle(
    aday_id: int,
    url: str,
    icerik: bytes,
    *,
    yayin_tarihi: str | None = None,
    baslik: str = "",
    mime_turu: str = "",
) -> dict[str, object]:
    if not resmi_meb_url(url):
        raise BelgeArsivHatasi("Yalnızca resmî meb.gov.tr belgeleri arşivlenebilir")
    uzanti = PurePosixPath(urlparse(url).path).suffix.lower()
    if uzanti not in DESTEKLENEN_UZANTILAR:
        raise BelgeArsivHatasi(f"Desteklenmeyen belge türü: {uzanti or 'bilinmiyor'}")
    if not icerik or len(icerik) > AZAMI_DOSYA_BOYUTU:
        raise BelgeArsivHatasi("Belge boş veya 20 MB sınırını aşıyor")

    ozet = sha256(icerik).hexdigest()
    try:
        yayin = date.fromisoformat(str(yayin_tarihi or ""))
    except ValueError:
        yayin = date.today()
    klasor = ARSIV_KOK / f"{yayin.year:04d}" / f"{yayin.month:02d}"
    klasor.mkdir(parents=True, exist_ok=True)
    hedef = klasor / _guvenli_dosya_adi(url, ozet)
    if not hedef.exists():
        gecici = hedef.with_suffix(hedef.suffix + ".tmp")
        gecici.write_bytes(icerik)
        gecici.replace(hedef)

    simdi = datetime.now().isoformat(timespec="seconds")
    metin = ""
    okuma_hatasi = None
    try:
        metin = belge_metni_oku(icerik, uzanti)
    except Exception as hata:  # Bozuk belge arşiv kaydını engellememeli.
        okuma_hatasi = str(hata)[:1000]
    veri = metin_verilerini_cikar(metin, baslik) if metin else {
        "belge_guveni": 0,
        "ham_metin": "",
    }
    kritik_tamam = (
        bool(veri.get("okul_adi"))
        and bool(veri.get("okul_turu"))
        and veri.get("ogrenci_sayisi") is not None
        and veri.get("muhammen_bedel_aylik") is not None
    )
    if uzanti == ".pdf" and metin and not kritik_tamam:
        try:
            derin_ocr = _pdf_ocr_metni(icerik, sayfa_siniri=24)
        except Exception:
            derin_ocr = ""
        if derin_ocr:
            metin = f"{metin}\n{derin_ocr}"
            veri = metin_verilerini_cikar(metin, baslik)
            kritik_tamam = (
                bool(veri.get("okul_adi"))
                and bool(veri.get("okul_turu"))
                and veri.get("ogrenci_sayisi") is not None
                and veri.get("muhammen_bedel_aylik") is not None
            )
    if kritik_tamam:
        durum = "analiz_edildi"
    elif metin:
        durum = "analiz_bekliyor"
        okuma_hatasi = (
            "Okul adı, okul türü, öğrenci sayısı ve aylık muhammen bedel yeniden işlenecek"
        )
    else:
        durum = "arsivlendi"
    mime = mime_turu or mimetypes.guess_type(str(hedef))[0] or "application/octet-stream"
    try:
        yerel_yol = str(hedef.relative_to(KOK))
    except ValueError:
        yerel_yol = str(hedef)
    with closing(baglan()) as conn, conn:
        cursor = conn.execute("""
            INSERT INTO ihale_belgeleri (
                aday_id, url, yerel_yol, sha256, boyut, mime_turu,
                durum, son_hata, ilk_indirme, son_kontrol
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(aday_id, url) DO UPDATE SET
                yerel_yol=excluded.yerel_yol, sha256=excluded.sha256,
                boyut=excluded.boyut, mime_turu=excluded.mime_turu,
                durum=excluded.durum, son_hata=excluded.son_hata,
                ilk_indirme=COALESCE(ihale_belgeleri.ilk_indirme, excluded.ilk_indirme),
                son_kontrol=excluded.son_kontrol
            RETURNING id
        """, (
            aday_id, url, yerel_yol, ozet, len(icerik), mime,
            durum, okuma_hatasi, simdi, simdi,
        ))
        belge_id = int(cursor.fetchone()[0])
        if metin:
            _analiz_verisini_yaz(conn, aday_id, belge_id, veri)
    return {
        "aday_id": aday_id,
        "belge_id": belge_id,
        "durum": durum,
        "yerel_yol": str(hedef),
        "sha256": ozet,
        "alan_sayisi": sum(
            veri.get(alan) is not None for alan in (
                "okul_adi", "okul_turu", "ogrenci_sayisi", "personel_sayisi",
                "muhammen_bedel_aylik", "muhammen_bedel_yillik",
                "sartname_bedeli", "gecici_teminat", "kantin_alani_m2",
            )
        ),
    }


def _hata_yaz(aday_id: int, url: str, hata: object) -> None:
    simdi = datetime.now().isoformat(timespec="seconds")
    with closing(baglan()) as conn, conn:
        conn.execute("""
            INSERT INTO ihale_belgeleri(aday_id, url, durum, son_hata, son_kontrol)
            VALUES (?, ?, 'hata', ?, ?)
            ON CONFLICT(aday_id, url) DO UPDATE SET
                durum='hata', son_hata=excluded.son_hata, son_kontrol=excluded.son_kontrol
        """, (aday_id, url, str(hata)[:1000], simdi))


def arsiv_kuyrugu(limit: int = 10) -> list[dict[str, object]]:
    """Tek okul ekini ana ilana, toplu ilan eklerini kendi okul kaydına bağlar."""
    with closing(baglan()) as conn:
        satirlar = conn.execute("""
            WITH hedefler AS (
                SELECT d.id,
                       CASE
                           WHEN d.eslesme_turu='ek_dosya'
                            AND parent.id IS NOT NULL
                            AND (
                                SELECT COUNT(*)
                                FROM duyuru_adaylari kardes
                                WHERE kardes.detay_url=d.detay_url
                                  AND kardes.eslesme_turu='ek_dosya'
                            )=1
                           THEN parent.id
                           ELSE d.id
                       END AS aday_id,
                       d.url, d.baslik, d.yayin_tarihi, d.dosya_turu,
                       d.eslesme_turu, d.durum
                FROM duyuru_adaylari d
                LEFT JOIN duyuru_adaylari parent
                  ON d.eslesme_turu='ek_dosya' AND parent.url=d.detay_url
            )
            SELECT h.aday_id, h.url, h.baslik, h.yayin_tarihi
            FROM hedefler h
            LEFT JOIN ihale_belgeleri b
              ON b.aday_id=h.aday_id AND b.url=h.url
            WHERE h.yayin_tarihi >= ?
              AND h.eslesme_turu IN ('dosya', 'toplu_dosya', 'ek_dosya')
              AND LOWER(COALESCE(h.dosya_turu, '')) IN (
                  '.pdf','.doc','.docx','.xls','.xlsx','.jpg','.jpeg','.png','.webp'
              )
              AND (
                    b.id IS NULL
                 OR (b.durum='hata' AND datetime(b.son_kontrol) <= datetime('now', '-6 hours'))
                 OR (b.durum='analiz_bekliyor'
                     AND datetime(b.son_kontrol) <= datetime('now', '-12 hours'))
            )
            ORDER BY CASE WHEN h.durum='aktif' THEN 0 ELSE 1 END,
                     CASE WHEN b.id IS NULL THEN 0 ELSE 1 END,
                     h.yayin_tarihi DESC, h.id DESC
            LIMIT ?
        """, (ihale_tarih_siniri().isoformat(), max(1, min(int(limit), 100))))
        return [dict(satir) for satir in satirlar]


def toplu_ilan_belgelerini_ayir(limit: int = 1000) -> dict[str, int]:
    """Yanlışlıkla ana ilana yazılan toplu ekleri indirmeden okul kayıtlarına taşır."""
    with closing(baglan()) as conn:
        satirlar = conn.execute("""
            WITH toplu_parent AS (
                SELECT parent.id AS parent_id
                FROM duyuru_adaylari parent
                JOIN duyuru_adaylari child
                  ON child.detay_url=parent.url
                 AND child.eslesme_turu='ek_dosya'
                GROUP BY parent.id
                HAVING COUNT(child.id) > 1
            )
            SELECT b.id AS eski_belge_id, b.aday_id AS parent_id,
                   b.url, b.yerel_yol, b.mime_turu,
                   child.id AS child_id, child.baslik, child.yayin_tarihi
            FROM ihale_belgeleri b
            JOIN toplu_parent tp ON tp.parent_id=b.aday_id
            JOIN duyuru_adaylari parent ON parent.id=tp.parent_id
            JOIN duyuru_adaylari child
              ON child.detay_url=parent.url
             AND child.eslesme_turu='ek_dosya'
             AND child.url=b.url
            ORDER BY b.id
            LIMIT ?
        """, (max(1, min(int(limit), 5000)),)).fetchall()
        parentler = [
            int(satir[0]) for satir in conn.execute("""
                SELECT parent.id
                FROM duyuru_adaylari parent
                JOIN duyuru_adaylari child
                  ON child.detay_url=parent.url
                 AND child.eslesme_turu='ek_dosya'
                GROUP BY parent.id
                HAVING COUNT(child.id) > 1
            """).fetchall()
        ]

    sonuc = {"ayrilan": 0, "eksik_dosya": 0, "hata": 0, "temizlenen_parent": 0}
    for satir in satirlar:
        yol = Path(str(satir["yerel_yol"] or ""))
        if not yol.is_absolute():
            yol = KOK / yol
        if not yol.is_file():
            sonuc["eksik_dosya"] += 1
            continue
        try:
            icerigi_arsivle(
                int(satir["child_id"]),
                str(satir["url"]),
                yol.read_bytes(),
                yayin_tarihi=str(satir["yayin_tarihi"] or ""),
                baslik=str(satir["baslik"] or ""),
                mime_turu=str(satir["mime_turu"] or ""),
            )
            with closing(baglan()) as conn, conn:
                conn.execute(
                    "DELETE FROM ihale_belgeleri WHERE id=?",
                    (int(satir["eski_belge_id"]),),
                )
            sonuc["ayrilan"] += 1
        except Exception:
            sonuc["hata"] += 1

    if parentler:
        yer_tutucular = ",".join("?" for _ in parentler)
        with closing(baglan()) as conn, conn:
            sonuc["temizlenen_parent"] = conn.execute(
                f"DELETE FROM ilan_analiz_verileri WHERE aday_id IN ({yer_tutucular})",
                parentler,
            ).rowcount
            conn.execute(
                f"DELETE FROM kantin_yatirim_analizleri WHERE aday_id IN ({yer_tutucular})",
                parentler,
            )
    return sonuc


def yerel_arsivi_yeniden_isle(
    limit: int = 5,
    *,
    bekleme_saati: int = 12,
) -> dict[str, int]:
    """Eksik zorunlu alanları indirme yapmadan yerel kopyadan yeniden çıkarır."""
    with closing(baglan()) as conn:
        satirlar = conn.execute("""
            SELECT b.aday_id, b.url, b.yerel_yol, b.mime_turu,
                   d.baslik, d.yayin_tarihi
            FROM ihale_belgeleri b
            JOIN duyuru_adaylari d ON d.id=b.aday_id
            LEFT JOIN ilan_analiz_verileri a ON a.aday_id=b.aday_id
            WHERE b.yerel_yol IS NOT NULL
              AND datetime(b.son_kontrol) <= datetime(
                    'now', printf('-%d hours', ?)
                  )
              AND (
                     a.aday_id IS NULL
                  OR NULLIF(TRIM(a.okul_adi), '') IS NULL
                  OR NULLIF(TRIM(a.okul_turu), '') IS NULL
                  OR a.ogrenci_sayisi IS NULL
                  OR a.muhammen_bedel_aylik IS NULL
              )
              AND NOT EXISTS (
                  SELECT 1
                  FROM duyuru_adaylari child
                  WHERE child.detay_url=d.url
                    AND child.eslesme_turu='ek_dosya'
                  GROUP BY child.detay_url
                  HAVING COUNT(*) > 1
              )
            ORDER BY CASE d.durum WHEN 'aktif' THEN 0 ELSE 1 END,
                     b.son_kontrol, b.id
            LIMIT ?
        """, (
            max(0, int(bekleme_saati)),
            max(1, min(int(limit), 500)),
        )).fetchall()
    sonuc = {"islenen": 0, "analiz_edilen": 0, "eksik_dosya": 0, "hata": 0}
    for satir in satirlar:
        yol = Path(str(satir["yerel_yol"] or ""))
        if not yol.is_absolute():
            yol = KOK / yol
        if not yol.is_file():
            sonuc["eksik_dosya"] += 1
            continue
        try:
            kayit = icerigi_arsivle(
                int(satir["aday_id"]),
                str(satir["url"]),
                yol.read_bytes(),
                yayin_tarihi=str(satir["yayin_tarihi"] or ""),
                baslik=str(satir["baslik"] or ""),
                mime_turu=str(satir["mime_turu"] or ""),
            )
            sonuc["islenen"] += 1
            sonuc["analiz_edilen"] += int(kayit["durum"] == "analiz_edildi")
        except Exception:
            sonuc["hata"] += 1
    return sonuc


def kayitli_metinleri_yeniden_ayristir(limit: int = 50) -> dict[str, int]:
    """OCR'ı tekrarlamadan saklanan metni güncel ayrıştırıcıyla yeniden işler."""
    with closing(baglan()) as conn:
        satirlar = conn.execute("""
            SELECT a.aday_id, a.kaynak_belge_id AS belge_id,
                   a.ham_metin, d.baslik, d.durum
            FROM ilan_analiz_verileri a
            JOIN duyuru_adaylari d ON d.id=a.aday_id
            JOIN ihale_belgeleri b ON b.id=a.kaynak_belge_id
            WHERE NULLIF(TRIM(a.ham_metin), '') IS NOT NULL
              AND (
                     NULLIF(TRIM(a.okul_adi), '') IS NULL
                  OR NULLIF(TRIM(a.okul_turu), '') IS NULL
                  OR a.ogrenci_sayisi IS NULL
                  OR a.muhammen_bedel_aylik IS NULL
              )
              AND NOT EXISTS (
                  SELECT 1
                  FROM duyuru_adaylari child
                  WHERE child.detay_url=d.url
                    AND child.eslesme_turu='ek_dosya'
                  GROUP BY child.detay_url
                  HAVING COUNT(*) > 1
              )
            ORDER BY CASE d.durum WHEN 'aktif' THEN 0 ELSE 1 END,
                     d.yayin_tarihi DESC, d.id DESC
            LIMIT ?
        """, (max(1, min(int(limit), 1000)),)).fetchall()

    sonuc = {"islenen": 0, "tamamlanan": 0, "eksik": 0, "hata": 0}
    for satir in satirlar:
        aday_id = int(satir["aday_id"])
        belge_id = int(satir["belge_id"])
        try:
            veri = metin_verilerini_cikar(
                str(satir["ham_metin"] or ""),
                str(satir["baslik"] or ""),
            )
            with closing(baglan()) as conn, conn:
                _analiz_verisini_yaz(conn, aday_id, belge_id, veri)
                birlesik = conn.execute("""
                    SELECT NULLIF(TRIM(okul_adi), '') IS NOT NULL,
                           NULLIF(TRIM(okul_turu), '') IS NOT NULL,
                           ogrenci_sayisi IS NOT NULL,
                           muhammen_bedel_aylik IS NOT NULL
                    FROM ilan_analiz_verileri
                    WHERE aday_id=?
                """, (aday_id,)).fetchone()
                tamam = bool(birlesik and all(bool(deger) for deger in birlesik))
                conn.execute("""
                    UPDATE ihale_belgeleri
                    SET durum=?, son_hata=?, son_kontrol=?
                    WHERE id=?
                """, (
                    "analiz_edildi" if tamam else "analiz_bekliyor",
                    None if tamam else (
                        "Zorunlu alanlar saklanan metinden yeniden ayrıştırılacak"
                    ),
                    datetime.now().isoformat(timespec="seconds"),
                    belge_id,
                ))
            sonuc["islenen"] += 1
            sonuc["tamamlanan" if tamam else "eksik"] += 1
        except Exception:
            sonuc["hata"] += 1
    return sonuc


def arsivi_geri_doldur(limit: int = 6) -> dict[str, int]:
    """Kuyruktaki küçük bir belge grubunu indirir; sonraki tur kaldığı yerden sürer."""
    kuyruk = arsiv_kuyrugu(limit)
    sonuc = {"islenen": 0, "analiz_edilen": 0, "hata": 0}
    if not kuyruk:
        _durum_yaz({**sonuc, "durum": "tamamlandi"})
        return sonuc
    oturum = requests.Session()
    oturum.headers.update(HEADERS)
    for kayit in kuyruk:
        aday_id = int(kayit["aday_id"])
        url = str(kayit["url"])
        try:
            if not resmi_meb_url(url):
                raise BelgeArsivHatasi("Resmî MEB alan adı değil")
            yanit = oturum.get(url, timeout=(8, 40), stream=True)
            yanit.raise_for_status()
            uzunluk = int(yanit.headers.get("content-length") or 0)
            if uzunluk > AZAMI_DOSYA_BOYUTU:
                raise BelgeArsivHatasi("Belge 20 MB sınırını aşıyor")
            parcalar: list[bytes] = []
            okunan = 0
            for parca in yanit.iter_content(chunk_size=128 * 1024):
                if not parca:
                    continue
                okunan += len(parca)
                if okunan > AZAMI_DOSYA_BOYUTU:
                    raise BelgeArsivHatasi("Belge 20 MB sınırını aşıyor")
                parcalar.append(parca)
            icerik = b"".join(parcalar)
            kaydedilen = icerigi_arsivle(
                aday_id,
                url,
                icerik,
                yayin_tarihi=str(kayit.get("yayin_tarihi") or ""),
                baslik=str(kayit.get("baslik") or ""),
                mime_turu=str(yanit.headers.get("content-type") or "").split(";", 1)[0],
            )
            sonuc["islenen"] += 1
            sonuc["analiz_edilen"] += int(kaydedilen["durum"] == "analiz_edildi")
        except Exception as hata:
            _hata_yaz(aday_id, url, hata)
            sonuc["hata"] += 1
    _durum_yaz({**sonuc, "durum": "calisiyor" if arsiv_kuyrugu(1) else "tamamlandi"})
    return sonuc


def arsiv_ozeti() -> dict[str, int]:
    with closing(baglan()) as conn:
        satir = conn.execute("""
            SELECT COUNT(*) AS toplam,
                   SUM(CASE WHEN durum='analiz_edildi' THEN 1 ELSE 0 END) AS analiz_edildi,
                   SUM(CASE WHEN durum='arsivlendi' THEN 1 ELSE 0 END) AS arsivlendi,
                   SUM(CASE WHEN durum='analiz_bekliyor' THEN 1 ELSE 0 END) AS analiz_bekliyor,
                   SUM(CASE WHEN durum='hata' THEN 1 ELSE 0 END) AS hata
            FROM ihale_belgeleri
        """).fetchone()
        bekleyen = conn.execute("""
            SELECT COUNT(*) FROM duyuru_adaylari d
            WHERE d.yayin_tarihi >= ?
              AND d.eslesme_turu IN ('dosya','toplu_dosya','ek_dosya')
              AND NOT EXISTS (
                  SELECT 1 FROM ihale_belgeleri b
                  WHERE b.url=d.url AND b.durum IN ('analiz_edildi','arsivlendi')
              )
        """, (ihale_tarih_siniri().isoformat(),)).fetchone()[0]
        aktif = conn.execute("""
            WITH hedefler AS (
                SELECT DISTINCT CASE
                    WHEN d.eslesme_turu='ek_dosya'
                     AND parent.id IS NOT NULL
                     AND (SELECT COUNT(*) FROM duyuru_adaylari kardes
                          WHERE kardes.detay_url=d.detay_url
                            AND kardes.eslesme_turu='ek_dosya')=1
                    THEN parent.id ELSE d.id END AS aday_id
                FROM duyuru_adaylari d
                LEFT JOIN duyuru_adaylari parent
                  ON d.eslesme_turu='ek_dosya' AND parent.url=d.detay_url
                WHERE d.durum='aktif'
                  AND d.eslesme_turu IN ('dosya','toplu_dosya','ek_dosya')
                  AND LOWER(COALESCE(d.dosya_turu, '')) IN (
                      '.pdf','.doc','.docx','.xls','.xlsx','.jpg','.jpeg','.png','.webp'
                  )
            )
            SELECT COUNT(*) AS toplam,
                   SUM(CASE WHEN EXISTS (
                       SELECT 1 FROM ilan_analiz_verileri a
                       WHERE a.aday_id=h.aday_id
                         AND NULLIF(TRIM(a.okul_adi), '') IS NOT NULL
                         AND NULLIF(TRIM(a.okul_turu), '') IS NOT NULL
                         AND a.ogrenci_sayisi IS NOT NULL
                         AND a.muhammen_bedel_aylik IS NOT NULL
                   ) THEN 1 ELSE 0 END) AS hazir
            FROM hedefler h
        """).fetchone()
    return {
        "toplam": int(satir["toplam"] or 0),
        "analiz_edildi": int(satir["analiz_edildi"] or 0),
        "arsivlendi": int(satir["arsivlendi"] or 0),
        "analiz_bekliyor": int(satir["analiz_bekliyor"] or 0),
        "hata": int(satir["hata"] or 0),
        "bekleyen": int(bekleyen or 0),
        "aktif_toplam": int(aktif["toplam"] or 0),
        "aktif_hazir": int(aktif["hazir"] or 0),
    }


def yerel_belgeyi_metadata_ile_aktar(
    metadata_dosyasi: str | Path,
    belge_dosyasi: str | Path,
) -> dict[str, object]:
    """Önceden indirilmiş resmî bir belgeyi metadata kaydıyla ana sisteme bağlar."""
    metadata = json.loads(Path(metadata_dosyasi).read_text(encoding="utf-8"))
    belge_yolu = Path(belge_dosyasi)
    url = str(metadata.get("sayfa_url") or "")
    if not resmi_meb_url(url):
        raise BelgeArsivHatasi("Metadata resmî MEB belge URL'si içermiyor")
    yayin_ham = str(metadata.get("yayin_tarihi") or "")
    yayin = None
    for format_ in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y"):
        try:
            yayin = datetime.strptime(yayin_ham, format_).date()
            break
        except ValueError:
            continue
    if yayin is None or yayin < ihale_tarih_siniri():
        raise BelgeArsivHatasi("Belge son bir yıllık yayın aralığında değil")
    ihale_ham = str(metadata.get("ihale_tarihi") or "")
    ihale = None
    for format_ in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y"):
        try:
            ihale = datetime.strptime(ihale_ham, format_).date().isoformat()
            break
        except ValueError:
            continue
    host = (urlparse(url).hostname or "").lower()
    simdi = datetime.now().isoformat(timespec="seconds")
    with closing(baglan()) as conn, conn:
        kaynak = conn.execute("""
            SELECT id FROM kaynaklar
            WHERE LOWER(url) LIKE ?
            ORDER BY CASE WHEN kaynak_seviyesi='il' THEN 0 ELSE 1 END, id
            LIMIT 1
        """, (f"https://{host}%",)).fetchone()
        if kaynak is None:
            raise BelgeArsivHatasi("Belgenin MEB kaynağı ana veritabanında bulunamadı")
        uzanti = belge_yolu.suffix.lower()
        conn.execute("""
            INSERT INTO duyuru_adaylari(
                kaynak_id, baslik, url, liste_url, eslesme_turu,
                yayin_tarihi, ihale_tarihi, dosya_turu,
                ilk_gorulme, son_gorulme
            ) VALUES (?, ?, ?, ?, 'toplu_dosya', ?, ?, ?, ?, ?)
            ON CONFLICT(url) DO UPDATE SET
                baslik=excluded.baslik, yayin_tarihi=excluded.yayin_tarihi,
                ihale_tarihi=COALESCE(duyuru_adaylari.ihale_tarihi, excluded.ihale_tarihi),
                son_gorulme=excluded.son_gorulme
        """, (
            int(kaynak["id"]),
            str(metadata.get("sayfa_basligi") or metadata.get("kaynak_basligi") or belge_yolu.stem),
            url,
            str(metadata.get("liste_sayfasi") or ""),
            yayin.isoformat(), ihale, uzanti, simdi, simdi,
        ))
        aday_id = int(conn.execute(
            "SELECT id FROM duyuru_adaylari WHERE url=?", (url,)
        ).fetchone()[0])
    sonuc = icerigi_arsivle(
        aday_id,
        url,
        belge_yolu.read_bytes(),
        yayin_tarihi=yayin.isoformat(),
        baslik=str(metadata.get("sayfa_basligi") or ""),
    )
    metin_dosyasi = None
    for okunan in metadata.get("okunan_dosyalar") or []:
        aday_metin = Path(str(okunan.get("metin_dosyasi") or ""))
        if aday_metin.is_file():
            metin_dosyasi = aday_metin
            break
    if metin_dosyasi is not None:
        veri = metin_verilerini_cikar(
            metin_dosyasi.read_text(encoding="utf-8"),
            str(metadata.get("sayfa_basligi") or ""),
        )
        with closing(baglan()) as conn, conn:
            _analiz_verisini_yaz(conn, aday_id, int(sonuc["belge_id"]), veri)
        sonuc["alan_sayisi"] = sum(
            veri.get(alan) is not None for alan in (
                "ogrenci_sayisi", "personel_sayisi", "muhammen_bedel",
                "sartname_bedeli", "gecici_teminat", "kantin_alani_m2",
            )
        )
    return sonuc


def _durum_yaz(degerler: dict[str, object]) -> None:
    veri = {**degerler, "guncelleme": datetime.now().isoformat(timespec="seconds")}
    gecici = DURUM_DOSYASI.with_suffix(".tmp")
    gecici.write_text(json.dumps(veri, ensure_ascii=False, indent=2), encoding="utf-8")
    gecici.replace(DURUM_DOSYASI)
