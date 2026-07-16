"""Resmî MEB sayfalarında yalnızca kantin ihale duyurularını keşfeder."""

from __future__ import annotations

from contextlib import closing
from datetime import date, datetime, timedelta
from functools import lru_cache
from io import BytesIO
from pathlib import PurePosixPath
from threading import Lock
from urllib.parse import urljoin, urlparse
import re
import unicodedata

import cv2
import numpy as np
import pypdfium2 as pdfium
import requests
from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from rapidocr_onnxruntime import RapidOCR

from veritabani import baglan, ihale_tarih_siniri

KANTIN_TERIMLERI = ("kantin", "büfe", "bufe", "çay ocağı", "cay ocagi", "kafeterya")
IHALE_TERIMLERI = ("ihale", "ilan", "duyuru", "kiralama", "kiraya", "şartname", "sartname")
GORSEL_UZANTILARI = {".jpg", ".jpeg", ".png", ".webp"}
DOSYA_UZANTILARI = {".pdf", ".doc", ".docx", ".xls", ".xlsx"} | GORSEL_UZANTILARI
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 Chrome/126.0 Safari/537.36 Ihalex/1.0"
    )
}
PDF_TARIH_TARAMA_GUNU = 120
AZAMI_IHALE_PLANLAMA_GUNU = 366
_OCR_KILIDI = Lock()


def _tarih_duzelt(deger: object) -> str | None:
    metin = str(deger or "").strip()
    eslesme = re.search(r"\b(\d{1,2}[/.]\d{1,2}[/.]\d{4}|\d{4}-\d{1,2}-\d{1,2})\b", metin)
    if not eslesme:
        return None
    tarih_metni = eslesme.group(1)
    for format_ in ("%d/%m/%Y", "%d.%m.%Y", "%Y-%m-%d"):
        try:
            tarih = datetime.strptime(tarih_metni, format_).date()
            return tarih.isoformat()
        except ValueError:
            continue
    return None


def _arama_metni(metin: str) -> str:
    sade = unicodedata.normalize("NFKD", str(metin).casefold()).replace("ı", "i")
    return "".join(harf for harf in sade if not unicodedata.combining(harf))


def _ihale_tarihi_makul_mu(aday: date, yayin: date | None) -> bool:
    """Kanun ve madde numaralarının tarih sanılmasını engeller."""
    if yayin:
        gun_farki = (aday - yayin).days
        return 0 < gun_farki <= AZAMI_IHALE_PLANLAMA_GUNU
    bugun = date.today()
    return bugun.year - 1 <= aday.year <= bugun.year + 1


def _ihale_tarihi_bul(metin: str, yayin_tarihi: str | None = None) -> str | None:
    """Yayın/mevzuat tarihlerini değil, ihale gününü bağlamıyla birlikte bulur."""
    arama = " ".join(_arama_metni(metin).split())
    tarih = r"(\d{1,2}[/.]\d{1,2}[/.]\d{4}|\d{4}-\d{1,2}-\d{1,2})"
    desenler = (
        rf"ihalenin\s+yapilacagi\s+tarih(?:\s*/\s*saat)?\s*[:：\-]?\s*{tarih}",
        rf"ihale\s+tarihi(?:\s+ve\s+saati|\s*/\s*saat)?\s*[:：\-]?\s*{tarih}",
        rf"ihale\s+gunu\s*[:：\-]?\s*{tarih}",
        rf"(?:[a-z]\s*\)\s*)?tarihi\s+ve\s+saati\s*[:：\-]?\s*{tarih}",
        rf"{tarih}\s*(?:tarihinde|gunu).{{0,80}}(?:ihale|kiralama)",
    )
    yayin = None
    if yayin_tarihi:
        try:
            yayin = date.fromisoformat(yayin_tarihi)
        except ValueError:
            pass
    bulunanlar: list[date] = []
    for desen in desenler:
        for eslesme in re.finditer(desen, arama, flags=re.IGNORECASE):
            duzeltilmis = _tarih_duzelt(eslesme.group(1))
            if not duzeltilmis:
                continue
            aday = date.fromisoformat(duzeltilmis)
            if not _ihale_tarihi_makul_mu(aday, yayin):
                continue
            bulunanlar.append(aday)
    if bulunanlar:
        return min(bulunanlar).isoformat()

    # Yeni belge şablonları için sabit başlığa bağlı olmayan bağlamsal yedek.
    puanli: list[tuple[int, date]] = []
    for eslesme in re.finditer(tarih, arama):
        duzeltilmis = _tarih_duzelt(eslesme.group(1))
        if not duzeltilmis:
            continue
        aday = date.fromisoformat(duzeltilmis)
        if not _ihale_tarihi_makul_mu(aday, yayin):
            continue
        cevre = arama[max(0, eslesme.start() - 140):eslesme.end() + 140]
        puan = 0
        puan += 6 if "ihale" in cevre else 0
        puan += 5 if "tarihi ve saati" in cevre or "yapilacagi" in cevre else 0
        puan += 3 if "saat" in cevre else 0
        puan += 2 if "gunu" in cevre else 0
        puan -= 8 if any(x in cevre for x in ("yonetmelik", "resmi gazete", "sayili")) else 0
        puan -= 4 if any(x in cevre for x in ("sozlesme", "baslama tarihi")) else 0
        if puan >= 6:
            puanli.append((puan, aday))
    if not puanli:
        return None
    en_yuksek = max(puan for puan, _ in puanli)
    return min(aday for puan, aday in puanli if puan == en_yuksek).isoformat()


@lru_cache(maxsize=1)
def _ocr_motoru() -> RapidOCR:
    return RapidOCR()


def _pdf_ocr_tarihi(
    icerik: bytes, yayin_tarihi: str | None, sayfa_siniri: int = 2
) -> str | None:
    """Görüntü tabanlı PDF'yi sayfa sayfa okur ve tarih bulununca durur."""
    pdf = pdfium.PdfDocument(BytesIO(icerik))
    for sayfa_no in range(min(len(pdf), sayfa_siniri)):
        goruntu = pdf[sayfa_no].render(scale=1.5).to_numpy()
        with _OCR_KILIDI:
            sonuc, _ = _ocr_motoru()(goruntu)
        if sonuc:
            metin = " ".join(str(satir[1]) for satir in sonuc)
            tarih = _ihale_tarihi_bul(metin, yayin_tarihi)
            if tarih:
                return tarih
    return None


def _gorsel_ocr_tarihi(icerik: bytes, yayin_tarihi: str | None) -> str | None:
    """JPG/PNG tabanlı ihale ilanını OCR ile okuyup gerçek ihale gününü bulur."""
    goruntu = cv2.imdecode(np.frombuffer(icerik, dtype=np.uint8), cv2.IMREAD_COLOR)
    if goruntu is None:
        return None
    with _OCR_KILIDI:
        sonuc, _ = _ocr_motoru()(goruntu)
    if not sonuc:
        return None
    metin = " ".join(str(satir[1]) for satir in sonuc)
    return _ihale_tarihi_bul(metin, yayin_tarihi)


def _gorsel_tarihi_getir(
    oturum: requests.Session, url: str, yayin_tarihi: str | None
) -> str | None:
    """Resmî MEB görsel ekinden ihale tarihini çıkarır."""
    try:
        if not yayin_tarihi:
            return None
        yanit = oturum.get(url, timeout=(10, 45))
        yanit.raise_for_status()
        if len(yanit.content) > 15_000_000:
            return None
        return _gorsel_ocr_tarihi(yanit.content, yayin_tarihi)
    except (OSError, ValueError, requests.RequestException):
        return None
    except Exception:
        # Bozuk veya desteklenmeyen görsel kaynak taramasını durdurmamalıdır.
        return None


def _docx_tarihi_getir(
    oturum: requests.Session, url: str, yayin_tarihi: str | None
) -> str | None:
    """Word DOCX ekindeki paragraf ve tabloları okuyup ihale tarihini çıkarır."""
    try:
        if not yayin_tarihi:
            return None
        yanit = oturum.get(url, timeout=(10, 45))
        yanit.raise_for_status()
        if len(yanit.content) > 15_000_000:
            return None
        belge = Document(BytesIO(yanit.content))
        parcalar = [paragraf.text for paragraf in belge.paragraphs if paragraf.text]
        for tablo in belge.tables:
            for satir in tablo.rows:
                parcalar.extend(hucre.text for hucre in satir.cells if hucre.text)
        return _ihale_tarihi_bul(" ".join(parcalar), yayin_tarihi)
    except (OSError, ValueError, requests.RequestException):
        return None
    except Exception:
        # Bozuk veya parola korumalı DOCX genel kaynak taramasını durdurmamalıdır.
        return None


def _xlsx_tarihi_getir(
    oturum: requests.Session, url: str, yayin_tarihi: str | None
) -> str | None:
    """XLSX ekindeki görünür hücreleri okuyup ihale tarihini çıkarır."""
    try:
        if not yayin_tarihi:
            return None
        yanit = oturum.get(url, timeout=(10, 45))
        yanit.raise_for_status()
        if len(yanit.content) > 15_000_000:
            return None
        kitap = load_workbook(BytesIO(yanit.content), read_only=True, data_only=True)
        parcalar: list[str] = []
        for sayfa in kitap.worksheets:
            for satir in sayfa.iter_rows(values_only=True):
                parcalar.extend(str(deger) for deger in satir if deger is not None)
        return _ihale_tarihi_bul(" ".join(parcalar), yayin_tarihi)
    except (OSError, ValueError, requests.RequestException):
        return None
    except Exception:
        return None


def _pdf_tarihi_getir(
    oturum: requests.Session,
    url: str,
    yayin_tarihi: str | None,
    gun_siniri: int | None = PDF_TARIH_TARAMA_GUNU,
    ocr_kullan: bool = True,
) -> str | None:
    """Yakın dönem PDF'lerinin ilk sayfalarından ihale tarihini çıkarır."""
    try:
        if not yayin_tarihi:
            return None
        yayin = date.fromisoformat(yayin_tarihi)
        if gun_siniri is not None and yayin < date.today() - timedelta(days=gun_siniri):
            return None
        yanit = oturum.get(url, timeout=(10, 45))
        yanit.raise_for_status()
        if len(yanit.content) > 15_000_000:
            return None
        okuyucu = PdfReader(BytesIO(yanit.content), strict=False)
        metin = " ".join((sayfa.extract_text() or "") for sayfa in okuyucu.pages[:6])
        ihale_tarihi = _ihale_tarihi_bul(metin, yayin_tarihi)
        if ihale_tarihi:
            return ihale_tarihi
        if not ocr_kullan:
            return None
        return _pdf_ocr_tarihi(yanit.content, yayin_tarihi)
    except (OSError, ValueError, requests.RequestException):
        return None
    except Exception:
        # Bozuk, taranmış veya şifreli PDF kaynak taramasını durdurmamalıdır.
        return None


def _tarih_gecerli(yayin_tarihi: str | None) -> bool:
    if not yayin_tarihi:
        return False
    try:
        return date.fromisoformat(yayin_tarihi) >= ihale_tarih_siniri()
    except ValueError:
        return False


def _resmi_meb_url(url: str) -> bool:
    host = (urlparse(url).hostname or "").lower()
    return host == "meb.gov.tr" or host.endswith(".meb.gov.tr")


def _yeni_ana_ilan_mi(eklenen_satir: int, eslesme_turu: str) -> bool:
    """Ek dosyalari yeni ihale sayacindan ve alarmindan ayirir."""
    return eklenen_satir == 1 and eslesme_turu != "ek_dosya"


def _ihale_eslesmesi(metin: str) -> bool:
    temiz = metin.casefold()
    return (any(k.casefold() in temiz for k in KANTIN_TERIMLERI)
            and any(k.casefold() in temiz for k in IHALE_TERIMLERI))


def _genel_baglanti_metni(metin: object) -> bool:
    temiz = unicodedata.normalize("NFKD", str(metin or "").casefold()).replace("ı", "i")
    temiz = "".join(harf for harf in temiz if not unicodedata.combining(harf))
    temiz = " ".join(re.findall(r"[a-z0-9]+", temiz))
    return not temiz or any(
        ifade in temiz for ifade in (
            "tiklayiniz", "tiklayin", "dosya icin", "duyuru icin",
            "ihale sartnamesi icin", "buradan ulasabilirsiniz",
        )
    )


def _html_getir(oturum: requests.Session, url: str) -> BeautifulSoup:
    yanit = oturum.get(url, timeout=(10, 30))
    yanit.raise_for_status()
    yanit.encoding = yanit.apparent_encoding or yanit.encoding
    return BeautifulSoup(yanit.text, "html.parser")


def _sayfa_adaylari(
    soup: BeautifulSoup, sayfa_url: str
) -> tuple[dict[str, tuple[str, str | None]], list[str]]:
    adaylar: dict[str, tuple[str, str | None]] = {}
    sayfalar: list[str] = []
    for baglanti in soup.select("a[href]"):
        tam_url = urljoin(sayfa_url, baglanti.get("href", "").strip())
        if not _resmi_meb_url(tam_url):
            continue
        baslik = " ".join(baglanti.get_text(" ", strip=True).split())
        cevre_metni = " ".join(baglanti.parent.get_text(" ", strip=True).split())
        yayin_tarihi = _tarih_duzelt(cevre_metni)
        if _ihale_eslesmesi(f"{baslik} {tam_url}"):
            if _tarih_gecerli(yayin_tarihi):
                adaylar[tam_url] = (
                    baslik or PurePosixPath(urlparse(tam_url).path).name,
                    yayin_tarihi,
                )
        if "/duyurular/kategori/2" in tam_url and tam_url != sayfa_url:
            sayfalar.append(tam_url)
    return adaylar, sayfalar


def _api_adaylari(
    oturum: requests.Session, kaynak: dict[str, object], sayfa_url: str
) -> dict[str, tuple[str, str | None]]:
    """MEB IYS DataTables uç noktasından render edilmemiş kayıtları alır."""
    strateji = str(kaynak.get("tarama_stratejisi") or "duyuru_listesi")
    parsed = urlparse(sayfa_url)
    origin = f"{parsed.scheme}://{parsed.netloc}"
    kategori_eslesme = re.search(r"/(?:kategori|dosya)/(\d+)", parsed.path)
    kategori = kategori_eslesme.group(1) if kategori_eslesme else "2"
    dosya_listesi = strateji == "toplu_dosya" or "/dosya/" in parsed.path
    endpoint = "/www/dosya_listele_ajax.php" if dosya_listesi else "/www/icerik_listele_ajax.php"
    headers = {"Referer": sayfa_url, "Origin": origin, "X-Requested-With": "XMLHttpRequest"}
    sonuc: dict[str, tuple[str, str | None]] = {}
    baslangic = 0
    while True:
        yanit = oturum.post(
            origin + endpoint,
            data={"draw": "1", "start": str(baslangic), "length": "1000",
                  "kategori": kategori, "dil": "tr"},
            headers=headers, timeout=(10, 30),
        )
        yanit.raise_for_status()
        veri = yanit.json()
        satirlar = veri.get("data") or []
        for satir in satirlar:
            baslik = " ".join(str(satir.get("BASLIK") or satir.get("ACIKLAMA") or "").split())
            tam_url = urljoin(sayfa_url, str(satir.get("LINK") or ""))
            yayin_tarihi = _tarih_duzelt(satir.get("ISLEMSAAT"))
            if not _tarih_gecerli(yayin_tarihi):
                continue
            if _resmi_meb_url(tam_url) and _ihale_eslesmesi(f"{baslik} {tam_url}"):
                sonuc[tam_url] = (baslik, yayin_tarihi)
        baslangic += len(satirlar)
        toplam = int(veri.get("recordsFiltered") or veri.get("recordsTotal") or 0)
        if not satirlar or baslangic >= toplam:
            break
    return sonuc


def _detay_dogrula(
    oturum: requests.Session,
    url: str,
    baslik: str,
    yayin_tarihi: str | None,
    dosya_gun_siniri: int | None = PDF_TARIH_TARAMA_GUNU,
    pdf_ocr_kullan: bool = True,
) -> tuple[bool, str, list[tuple[str, str]], str | None]:
    uzanti = PurePosixPath(urlparse(url).path).suffix.lower()
    if uzanti in DOSYA_UZANTILARI:
        ihale_tarihi = None
        if uzanti == ".pdf":
            ihale_tarihi = _pdf_tarihi_getir(
                oturum, url, yayin_tarihi, gun_siniri=dosya_gun_siniri,
                ocr_kullan=pdf_ocr_kullan,
            )
        elif uzanti in GORSEL_UZANTILARI:
            ihale_tarihi = _gorsel_tarihi_getir(oturum, url, yayin_tarihi)
        elif uzanti == ".docx":
            ihale_tarihi = _docx_tarihi_getir(oturum, url, yayin_tarihi)
        elif uzanti == ".xlsx":
            ihale_tarihi = _xlsx_tarihi_getir(oturum, url, yayin_tarihi)
        return _ihale_eslesmesi(f"{baslik} {url}"), "dosya", [], ihale_tarihi
    soup = _html_getir(oturum, url)
    ana_metin = " ".join(soup.get_text(" ", strip=True).split())
    dogru = _ihale_eslesmesi(f"{baslik} {url} {ana_metin[:12000]}")
    ihale_tarihi = _ihale_tarihi_bul(ana_metin, yayin_tarihi)
    ekler: list[tuple[str, str]] = []
    if dogru:
        for baglanti in soup.select("a[href]"):
            ek_url = urljoin(url, baglanti.get("href", "").strip())
            ek_uzanti = PurePosixPath(urlparse(ek_url).path).suffix.lower()
            if _resmi_meb_url(ek_url) and ek_uzanti in DOSYA_UZANTILARI:
                ek_baslik = " ".join(baglanti.get_text(" ", strip=True).split())
                if _genel_baglanti_metni(ek_baslik):
                    ek_baslik = baslik
                ekler.append((ek_url, ek_baslik))
        if not ihale_tarihi:
            for ek_url, _ in ekler:
                ek_uzanti = PurePosixPath(urlparse(ek_url).path).suffix.lower()
                if ek_uzanti == ".pdf":
                    ihale_tarihi = _pdf_tarihi_getir(
                        oturum, ek_url, yayin_tarihi, gun_siniri=dosya_gun_siniri,
                        ocr_kullan=pdf_ocr_kullan,
                    )
                elif ek_uzanti in GORSEL_UZANTILARI:
                    ihale_tarihi = _gorsel_tarihi_getir(oturum, ek_url, yayin_tarihi)
                elif ek_uzanti == ".docx":
                    ihale_tarihi = _docx_tarihi_getir(oturum, ek_url, yayin_tarihi)
                elif ek_uzanti == ".xlsx":
                    ihale_tarihi = _xlsx_tarihi_getir(oturum, ek_url, yayin_tarihi)
                if ihale_tarihi:
                    break
    return dogru, "detay", ekler, ihale_tarihi


def _ham_adaylari_kaydet(
    kaynak_id: int,
    liste_url: str,
    adaylar: dict[str, tuple[str, str | None]],
) -> int:
    """Keşfedilen metadatayı detay bağlantısından bağımsız olarak kalıcılaştırır."""
    simdi = datetime.now().isoformat(timespec="seconds")
    yazilan = 0
    with closing(baglan()) as conn, conn:
        for aday_url, (baslik, yayin_tarihi) in adaylar.items():
            if not _tarih_gecerli(yayin_tarihi):
                continue
            cursor = conn.execute("""
                INSERT INTO ham_duyurular (
                    kaynak_id, baslik, url, liste_url, yayin_tarihi,
                    ilk_gorulme, son_gorulme
                ) VALUES (?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(url) DO UPDATE SET
                    baslik=excluded.baslik,
                    liste_url=excluded.liste_url,
                    yayin_tarihi=excluded.yayin_tarihi,
                    son_gorulme=excluded.son_gorulme
            """, (
                kaynak_id, baslik, aday_url, liste_url, yayin_tarihi, simdi, simdi
            ))
            yazilan += cursor.rowcount
    return yazilan


def _ham_aday_sonucu_yaz(
    aday_url: str,
    *,
    erisim_durumu: str,
    dogrulama_durumu: str,
    hata: object | None = None,
) -> None:
    with closing(baglan()) as conn, conn:
        conn.execute("""
            UPDATE ham_duyurular
            SET erisim_durumu=?, dogrulama_durumu=?,
                deneme_sayisi=deneme_sayisi+1, son_hata=?
            WHERE url=?
        """, (
            erisim_durumu,
            dogrulama_durumu,
            str(hata)[:1000] if hata else None,
            aday_url,
        ))


def kaynak_tara(kaynak: dict[str, object]) -> tuple[int, int]:
    url = str(kaynak["url"])
    if not _resmi_meb_url(url):
        raise ValueError("Yalnızca resmî meb.gov.tr kaynakları taranabilir")
    baslangic = datetime.now().isoformat(timespec="seconds")
    ilk_tarama = not bool(kaynak.get("son_basarili_tarama"))
    with closing(baglan()) as conn, conn:
        tarama_id = conn.execute("""
            INSERT INTO kaynak_taramalari(kaynak_id, baslangic, durum)
            VALUES (?, ?, 'tariyor')
        """, (kaynak["id"], baslangic)).lastrowid

    try:
        oturum = requests.Session()
        oturum.headers.update(HEADERS)
        ilk_soup = _html_getir(oturum, url)
        adaylar, sayfalar = _sayfa_adaylari(ilk_soup, url)
        try:
            adaylar.update(_api_adaylari(oturum, kaynak, url))
        except (requests.RequestException, ValueError, KeyError):
            # Eski MEB şablonlarında API yoksa görünür HTML bağlantıları kullanılır.
            if kaynak.get("tarama_stratejisi") == "duyuru_listesi":
                for sayfa_url in list(dict.fromkeys(sayfalar))[:4]:
                    try:
                        ek_adaylar, _ = _sayfa_adaylari(
                            _html_getir(oturum, sayfa_url), sayfa_url
                        )
                        adaylar.update(ek_adaylar)
                    except requests.RequestException:
                        continue

        _ham_adaylari_kaydet(int(kaynak["id"]), url, adaylar)

        dogrulanmis: list[
            tuple[str, str, str, str | None, str | None, str | None]
        ] = []
        for aday_url, (baslik, yayin_tarihi) in adaylar.items():
            if not _tarih_gecerli(yayin_tarihi):
                continue
            if kaynak.get("tarama_stratejisi") == "toplu_dosya":
                ihale_tarihi = _ihale_tarihi_bul(baslik, yayin_tarihi)
                if PurePosixPath(urlparse(aday_url).path).suffix.lower() == ".pdf":
                    ihale_tarihi = ihale_tarihi or _pdf_tarihi_getir(
                        oturum, aday_url, yayin_tarihi
                    )
                dogrulanmis.append(
                    (aday_url, baslik, "toplu_dosya", None, yayin_tarihi, ihale_tarihi)
                )
                _ham_aday_sonucu_yaz(
                    aday_url, erisim_durumu="eristi", dogrulama_durumu="dogrulandi"
                )
                continue
            try:
                dogru, eslesme_turu, ekler, ihale_tarihi = _detay_dogrula(
                    oturum, aday_url, baslik, yayin_tarihi
                )
            except requests.RequestException as hata:
                _ham_aday_sonucu_yaz(
                    aday_url,
                    erisim_durumu="hata",
                    dogrulama_durumu="bekliyor",
                    hata=hata,
                )
                continue
            if not dogru:
                _ham_aday_sonucu_yaz(
                    aday_url, erisim_durumu="eristi", dogrulama_durumu="reddedildi"
                )
                continue
            _ham_aday_sonucu_yaz(
                aday_url, erisim_durumu="eristi", dogrulama_durumu="dogrulandi"
            )
            dogrulanmis.append(
                (aday_url, baslik, eslesme_turu, aday_url, yayin_tarihi, ihale_tarihi)
            )
            for ek_url, ek_baslik in ekler:
                dogrulanmis.append(
                    (ek_url, ek_baslik, "ek_dosya", aday_url, yayin_tarihi, ihale_tarihi)
                )

        simdi = datetime.now().isoformat(timespec="seconds")
        yeni = 0
        with closing(baglan()) as conn, conn:
            for (
                aday_url, baslik, eslesme_turu, detay_url, yayin_tarihi, ihale_tarihi
            ) in dogrulanmis:
                uzanti = PurePosixPath(urlparse(aday_url).path).suffix.lower()
                cursor = conn.execute("""
                    INSERT OR IGNORE INTO duyuru_adaylari (
                        kaynak_id, baslik, url, detay_url, liste_url,
                        eslesme_turu, yayin_tarihi, ihale_tarihi, dosya_turu,
                        ilk_gorulme, son_gorulme
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (kaynak["id"], baslik, aday_url, detay_url, url,
                      eslesme_turu, yayin_tarihi, ihale_tarihi,
                      uzanti if uzanti in DOSYA_UZANTILARI else "html",
                      simdi, simdi))
                yeni_ana_ilan = _yeni_ana_ilan_mi(cursor.rowcount, eslesme_turu)
                if yeni_ana_ilan:
                    yeni += 1
                if yeni_ana_ilan and not ilk_tarama:
                    conn.execute("""
                        INSERT OR IGNORE INTO alarmlar(aday_id, kanal, olusturma_tarihi)
                        VALUES (?, 'telegram', ?)
                    """, (cursor.lastrowid, simdi))
                conn.execute("""
                    UPDATE duyuru_adaylari
                    SET son_gorulme=?, ihale_tarihi=COALESCE(ihale_tarihi, ?)
                    WHERE url=?
                """, (simdi, ihale_tarihi, aday_url))
            conn.execute("""
                UPDATE kaynaklar SET son_tarama=?, son_basarili_tarama=?,
                    son_durum='basarili', son_hata=NULL WHERE id=?
            """, (simdi, simdi, kaynak["id"]))
            conn.execute("""
                UPDATE kaynak_taramalari SET bitis=?, durum='basarili',
                    bulunan_baglanti=?, yeni_baglanti=? WHERE id=?
            """, (simdi, len(dogrulanmis), yeni, tarama_id))
        return len(dogrulanmis), yeni
    except Exception as hata:
        simdi = datetime.now().isoformat(timespec="seconds")
        mesaj = str(hata)[:1000]
        with closing(baglan()) as conn, conn:
            conn.execute("UPDATE kaynaklar SET son_tarama=?, son_durum='hata', son_hata=? WHERE id=?",
                         (simdi, mesaj, kaynak["id"]))
            conn.execute("UPDATE kaynak_taramalari SET bitis=?, durum='hata', hata=? WHERE id=?",
                         (simdi, mesaj, tarama_id))
        raise
