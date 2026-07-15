import json
import os
import re
from pathlib import Path

import pytesseract
from docx import Document
from pdf2image import convert_from_path
from pypdf import PdfReader


GIRDI_DOSYASI = "meb_icerik_sonucu.json"

CIKTI_DOSYASI = "meb_ihale_verisi.json"

METIN_KLASORU = (
    Path("ilanlar")
    /
    "meb_metinleri"
)

TESSERACT_YOLU = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)

POPPLER_YOLU = (
    r"C:\poppler\Library\bin"
)


pytesseract.pytesseract.tesseract_cmd = (
    TESSERACT_YOLU
)


def metin_temizle(metin):

    if not metin:

        return ""


    metin = metin.replace(
        "\xa0",
        " "
    )

    metin = metin.replace(
        "\r",
        "\n"
    )

    metin = metin.replace(
        "m²",
        " m2 "
    )

    metin = metin.replace(
        "M²",
        " m2 "
    )

    metin = re.sub(
        r"[ \t]+",
        " ",
        metin
    )

    metin = re.sub(
        r"\n{3,}",
        "\n\n",
        metin
    )


    return metin.strip()


def turkce_normalize(metin):

    if not metin:

        return ""


    ceviri = str.maketrans(
        {
            "ç": "c",
            "Ç": "c",
            "ğ": "g",
            "Ğ": "g",
            "ı": "i",
            "İ": "i",
            "ö": "o",
            "Ö": "o",
            "ş": "s",
            "Ş": "s",
            "ü": "u",
            "Ü": "u"
        }
    )


    return (
        str(
            metin
        )
        .translate(
            ceviri
        )
        .lower()
    )


def metin_kaliteli_mi(metin):

    if not metin:

        return False


    temiz = metin_temizle(
        metin
    )


    if len(
        temiz
    ) < 100:

        return False


    harfler = re.findall(
        r"[A-Za-zÇĞİÖŞÜçğıöşü]",
        temiz
    )


    sayilar = re.findall(
        r"\d",
        temiz
    )


    toplam = max(
        len(
            temiz
        ),
        1
    )


    harf_orani = (
        len(
            harfler
        )
        /
        toplam
    )


    sayi_orani = (
        len(
            sayilar
        )
        /
        toplam
    )


    kelimeler = re.findall(
        r"[A-Za-zÇĞİÖŞÜçğıöşü]{3,}",
        temiz
    )


    benzersiz_kelimeler = set(
        kelime.casefold()
        for kelime in kelimeler
    )


    uzun_sayilar = re.findall(
        r"\b\d{6,}\b",
        temiz
    )


    benzersiz_uzun_sayilar = set(
        uzun_sayilar
    )


    if harf_orani < 0.12:

        return False


    if sayi_orani > 0.70:

        return False


    if len(
        benzersiz_kelimeler
    ) < 8:

        return False


    if (
        len(
            uzun_sayilar
        ) >= 5
        and
        len(
            benzersiz_uzun_sayilar
        ) <= 2
    ):

        return False


    return True


def metin_dosyasi_kaydet(
        kaynak_dosya,
        metin
):

    METIN_KLASORU.mkdir(
        parents=True,
        exist_ok=True
    )


    hedef = (
        METIN_KLASORU
        /
        (
            Path(
                kaynak_dosya
            ).stem
            + ".txt"
        )
    )


    with open(
        hedef,
        "w",
        encoding="utf-8"
    ) as dosya:

        dosya.write(
            metin
        )


    return str(
        hedef
    )


def pdf_normal_oku(dosya_yolu):

    metinler = []


    try:

        pdf = PdfReader(
            dosya_yolu
        )


        for sayfa in pdf.pages:

            sayfa_metni = (
                sayfa.extract_text()
                or ""
            )


            if sayfa_metni.strip():

                metinler.append(
                    sayfa_metni
                )


    except Exception as hata:

        print(
            "PDF normal okuma hatası:"
        )

        print(hata)

        return ""


    return metin_temizle(
        "\n".join(
            metinler
        )
    )


def pdf_ocr_oku(dosya_yolu):

    print(
        "PDF OCR ile okunuyor..."
    )


    try:

        sayfalar = convert_from_path(
            dosya_yolu,
            dpi=300,
            poppler_path=POPPLER_YOLU
        )


    except Exception as hata:

        print(
            "PDF görüntü dönüştürme hatası:"
        )

        print(hata)

        return ""


    metinler = []


    for sayfa_no, sayfa in enumerate(
        sayfalar,
        start=1
    ):

        print(
            f"  {sayfa_no}. sayfa OCR okunuyor..."
        )


        try:

            sayfa_metni = (
                pytesseract.image_to_string(
                    sayfa,
                    lang="tur",
                    config="--psm 6"
                )
            )


        except Exception as hata:

            print(
                "Türkçe OCR hatası:"
            )

            print(hata)

            print(
                "İngilizce OCR deneniyor..."
            )


            try:

                sayfa_metni = (
                    pytesseract.image_to_string(
                        sayfa,
                        lang="eng",
                        config="--psm 6"
                    )
                )


            except Exception as ikinci_hata:

                print(
                    "OCR tamamen başarısız:"
                )

                print(
                    ikinci_hata
                )

                sayfa_metni = ""


        if sayfa_metni.strip():

            metinler.append(
                sayfa_metni
            )


    return metin_temizle(
        "\n".join(
            metinler
        )
    )


def pdf_oku(dosya_yolu):

    normal_metin = pdf_normal_oku(
        dosya_yolu
    )


    if metin_kaliteli_mi(
        normal_metin
    ):

        print(
            "PDF kaliteli metin katmanından okundu."
        )

        return normal_metin


    if normal_metin:

        print(
            "PDF metin katmanı bulundu ancak geçersiz."
        )

        print(
            "Metin örneği:"
        )

        print(
            normal_metin[:200]
        )


    else:

        print(
            "PDF metin katmanı boş."
        )


    ocr_metni = pdf_ocr_oku(
        dosya_yolu
    )


    if not ocr_metni:

        print(
            "OCR sonucunda metin çıkarılamadı."
        )

        return ""


    if metin_kaliteli_mi(
        ocr_metni
    ):

        print(
            "PDF OCR ile başarıyla okundu."
        )


    else:

        print(
            "UYARI: OCR metni düşük kaliteli olabilir."
        )


    return ocr_metni


def docx_oku(dosya_yolu):

    metinler = []


    try:

        belge = Document(
            dosya_yolu
        )


        for paragraf in belge.paragraphs:

            yazi = metin_temizle(
                paragraf.text
            )


            if yazi:

                metinler.append(
                    yazi
                )


        for tablo_no, tablo in enumerate(
            belge.tables,
            start=1
        ):

            metinler.append(
                f"TABLO {tablo_no}"
            )


            for satir in tablo.rows:

                hucreler = []


                for hucre in satir.cells:

                    yazi = metin_temizle(
                        hucre.text
                    )


                    if yazi:

                        hucreler.append(
                            yazi
                        )


                if hucreler:

                    metinler.append(
                        " | ".join(
                            hucreler
                        )
                    )


    except Exception as hata:

        print(
            "DOCX okuma hatası:"
        )

        print(hata)

        return ""


    return metin_temizle(
        "\n".join(
            metinler
        )
    )


def ek_dosya_oku(dosya_yolu):

    uzanti = Path(
        dosya_yolu
    ).suffix.lower()


    if uzanti == ".pdf":

        return pdf_oku(
            dosya_yolu
        )


    if uzanti == ".docx":

        return docx_oku(
            dosya_yolu
        )


    return ""


def para_temizle(deger):

    if not deger:

        return None


    metin = str(
        deger
    )


    metin = re.sub(
        r"[^\d.,]",
        "",
        metin
    )


    if not metin:

        return None


    if (
        "." in metin
        and
        "," in metin
    ):

        metin = metin.replace(
            ".",
            ""
        )

        metin = metin.replace(
            ",",
            "."
        )


    elif "," in metin:

        metin = metin.replace(
            ",",
            "."
        )


    elif "." in metin:

        parcalar = metin.split(
            "."
        )


        if (
            len(
                parcalar
            ) > 1
            and
            all(
                len(
                    parca
                ) == 3
                for parca in parcalar[1:]
            )
        ):

            metin = metin.replace(
                ".",
                ""
            )


    try:

        sayi = float(
            metin
        )


        if sayi.is_integer():

            return int(
                sayi
            )


        return sayi


    except ValueError:

        return None


def okul_adi_bul(metin):

    desenler = [

        (
            r"OKULUN\s+ADI"
            r"\s*[|:\-]?\s*"
            r"([^\n|]{4,150}?"
            r"(?:İlkokulu|Ortaokulu|"
            r"Anadolu Lisesi|Lisesi))"
        ),

        (
            r"([A-ZÇĞİÖŞÜ]"
            r"[A-ZÇĞİÖŞÜa-zçğıöşü0-9 "
            r"\-.'’]{3,140}?"
            r"(?:İlkokulu|Ortaokulu|"
            r"Anadolu Lisesi|Lisesi))"
        )

    ]


    for desen in desenler:

        sonuc = re.search(
            desen,
            metin,
            re.IGNORECASE
        )


        if sonuc:

            okul = re.sub(
                r"\s+",
                " ",
                sonuc.group(
                    1
                )
            )


            return okul.strip(
                " \"'|-:"
            )


    return ""


def ilce_bul(metin):

    desenler = [

        (
            r"([A-ZÇĞİÖŞÜa-zçğıöşü "
            r"\-]+?)\s+İLÇESİ"
        ),

        (
            r"BULUNDUĞU\s+İLÇE"
            r"\s*[|:\-]?\s*"
            r"([^\n|]{3,80})"
        )

    ]


    for desen in desenler:

        sonuc = re.search(
            desen,
            metin,
            re.IGNORECASE
        )


        if sonuc:

            ilce = re.sub(
                r"\s+",
                " ",
                sonuc.group(
                    1
                )
            )


            return ilce.strip(
                " |-:"
            ).title()


    return ""


def ogrenci_sayisi_bul(metin):

    desenler = [

        (
            r"ÖĞRENCİ\s+SAYISI"
            r"(?:\s+VE\s+ÖĞRENİM\s+DURUMU)?"
            r"\s*[|:\-]?\s*"
            r"(\d{2,5})"
        ),

        (
            r"ÖĞRENCİ\s+MEVCUDU"
            r"\s*[|:\-]?\s*"
            r"(\d{2,5})"
        ),

        (
            r"(\d{2,5})"
            r"\s+ÖĞRENCİ"
        )

    ]


    for desen in desenler:

        sonuc = re.search(
            desen,
            metin,
            re.IGNORECASE
        )


        if sonuc:

            sayi = int(
                sonuc.group(
                    1
                )
            )


            if 50 <= sayi <= 20000:

                return sayi


    return None


def personel_sayisi_bul(metin):

    desenler = [

        (
            r"ÖĞRETMEN\s+VE\s+DİĞER\s+"
            r"PERSONEL\s+SAYISI"
            r"\s*[|:\-]?\s*"
            r"(\d{1,4})"
        ),

        (
            r"PERSONEL\s+SAYISI"
            r"\s*[|:\-]?\s*"
            r"(\d{1,4})"
        )

    ]


    for desen in desenler:

        sonuc = re.search(
            desen,
            metin,
            re.IGNORECASE
        )


        if sonuc:

            sayi = int(
                sonuc.group(
                    1
                )
            )


            if 1 <= sayi <= 3000:

                return sayi


    return None


def alan_bul(metin):

    desenler = [

        (
            r"KANTİN\s+YERİ\s+VE\s+"
            r"M[²2?'’`]*\s*[']?[SŞ]İ"
            r"\s*[|:\-]?\s*"
            r"([\d.,]+)"
        ),

        (
            r"KANTİN\s+ALANI"
            r"\s*[|:\-]?\s*"
            r"([\d.,]+)"
            r"\s*(?:M2|M²)"
        ),

        (
            r"([\d.,]+)"
            r"\s*(?:M2|M²)"
        )

    ]


    for desen in desenler:

        sonuc = re.search(
            desen,
            metin,
            re.IGNORECASE
        )


        if sonuc:

            alan = para_temizle(
                sonuc.group(
                    1
                )
            )


            if (
                alan is not None
                and
                2 <= alan <= 5000
            ):

                return alan


    return None


def muhammen_bedel_bul(metin):

    desenler = [

        (
            r"MUHAMMEN\s+BEDEL"
            r"\s*[|:\-]?\s*"
            r"([\d.,]+)"
            r"(?:\s*\([^)]*\))?"
            r"\s*[-–]?\s*TL"
        ),

        (
            r"AYLIK\s+MUHAMMEN\s+"
            r"(?:KİRA\s+)?BEDELİ"
            r"\s*[|:\-]?\s*"
            r"([\d.,]+)"
        )

    ]


    for desen in desenler:

        sonuc = re.search(
            desen,
            metin,
            re.IGNORECASE
        )


        if sonuc:

            return para_temizle(
                sonuc.group(
                    1
                )
            )


    return None


def gecici_teminat_bul(metin):

    desenler = [

        (
            r"GEÇİCİ\s+TEMİNAT"
            r".{0,250}?"
            r"(?:EN\s+AZ\s*)?"
            r"([\d.]+,\d{2})"
            r"\s*TL"
        ),

        (
            r"GEÇİCİ\s+TEMİNAT"
            r".{0,250}?"
            r"([\d.]+)"
            r"\s*TL"
        ),

        (
            r"EN\s+AZ\s+"
            r"([\d.]+,\d{2})"
            r"\s*TL"
            r".{0,80}?"
            r"GEÇİCİ\s+TEMİNAT"
        )

    ]


    for desen in desenler:

        sonuc = re.search(
            desen,
            metin,
            re.IGNORECASE
            |
            re.DOTALL
        )


        if sonuc:

            tutar = para_temizle(
                sonuc.group(
                    1
                )
            )


            if (
                tutar is not None
                and
                tutar >= 100
            ):

                return tutar


    return None


def ihale_tarihi_saati_bul(metin):

    desenler = [

        (
            r"İHALENİN\s+TARİHİ\s+VE\s+SAATİ"
            r"\s*[|:\-]?\s*"
            r"(\d{1,2}[./]\d{1,2}[./]\d{4})"
            r"\s*[-–—]?\s*"
            r"([0-2]?\d[:.][0-5]\d)"
        ),

        (
            r"İHALE\s+TARİHİ\s+VE\s+SAATİ"
            r"\s*[|:\-]?\s*"
            r"(\d{1,2}[./]\d{1,2}[./]\d{4})"
            r"\s*[-–—]?\s*"
            r"([0-2]?\d[:.][0-5]\d)"
        ),

        (
            r"(\d{1,2}[./]\d{1,2}[./]\d{4})"
            r"\s*[-–—]\s*"
            r"([0-2]?\d[:.][0-5]\d)"
        )

    ]


    for desen in desenler:

        sonuc = re.search(
            desen,
            metin,
            re.IGNORECASE
        )


        if sonuc:

            tarih = sonuc.group(
                1
            ).replace(
                "/",
                "."
            )


            saat = sonuc.group(
                2
            ).replace(
                ".",
                ":"
            )


            return tarih, saat


    return "", ""


def yayin_araligi_bul(metin):

    desen = (
        r"BU\s+İLAN"
        r".{0,100}?"
        r"(\d{1,2}[./]\d{1,2}[./]\d{4})"
        r".{0,60}?"
        r"SAAT\s+([0-2]?\d[:.][0-5]\d)"
        r".{0,200}?"
        r"(\d{1,2}[./]\d{1,2}[./]\d{4})"
        r".{0,60}?"
        r"SAAT\s+([0-2]?\d[:.][0-5]\d)"
    )


    sonuc = re.search(
        desen,
        metin,
        re.IGNORECASE
        |
        re.DOTALL
    )


    if not sonuc:

        return {
            "ilan_baslangic_tarihi": "",
            "ilan_baslangic_saati": "",
            "ilan_bitis_tarihi": "",
            "ilan_bitis_saati": ""
        }


    return {

        "ilan_baslangic_tarihi":
            sonuc.group(
                1
            ).replace(
                "/",
                "."
            ),

        "ilan_baslangic_saati":
            sonuc.group(
                2
            ).replace(
                ".",
                ":"
            ),

        "ilan_bitis_tarihi":
            sonuc.group(
                3
            ).replace(
                "/",
                "."
            ),

        "ilan_bitis_saati":
            sonuc.group(
                4
            ).replace(
                ".",
                ":"
            )

    }


def tablo_satiri_bul(metin):

    desen = (
        r"([A-ZÇĞİÖŞÜa-zçğıöşü0-9 "
        r"\-.'’]+(?:İlkokulu|Ortaokulu|"
        r"Anadolu Lisesi|Lisesi))"
        r"\s*\|\s*"
        r"(Kantin)"
        r"\s*\|\s*"
        r"(\d{2,5})"
        r"\s*\|\s*"
        r"([\d.,]+)\s*TL"
        r"\s*\|\s*"
        r"([\d.,]+)"
        r"(?:\s*TL)?"
        r"\s*\|\s*"
        r"(\d{1,2}[./]\d{1,2}[./]\d{4})"
        r"\s*\|\s*"
        r"(\d{1,2}[:.]\d{2})"
    )


    sonuc = re.search(
        desen,
        metin,
        re.IGNORECASE
    )


    if not sonuc:

        return None


    return {

        "okul_adi":
            re.sub(
                r"\s+",
                " ",
                sonuc.group(
                    1
                )
            ).strip(),

        "ilce":
            "",

        "is_yeri_turu":
            sonuc.group(
                2
            ).strip(),

        "ogrenci_sayisi":
            int(
                sonuc.group(
                    3
                )
            ),

        "personel_sayisi":
            None,

        "alan_m2":
            None,

        "aylik_kira":
            para_temizle(
                sonuc.group(
                    4
                )
            ),

        "gecici_teminat":
            para_temizle(
                sonuc.group(
                    5
                )
            ),

        "ihale_tarihi":
            sonuc.group(
                6
            ).replace(
                "/",
                "."
            ),

        "ihale_saati":
            sonuc.group(
                7
            ).replace(
                ".",
                ":"
            ),

        "veri_yontemi":
            "tablo_satiri"

    }


def genel_belge_verisi_cikar(metin):

    ihale_tarihi, ihale_saati = (
        ihale_tarihi_saati_bul(
            metin
        )
    )


    yayin_araligi = yayin_araligi_bul(
        metin
    )


    sonuc = {

        "okul_adi":
            okul_adi_bul(
                metin
            ),

        "ilce":
            ilce_bul(
                metin
            ),

        "is_yeri_turu":
            "Kantin",

        "ogrenci_sayisi":
            ogrenci_sayisi_bul(
                metin
            ),

        "personel_sayisi":
            personel_sayisi_bul(
                metin
            ),

        "alan_m2":
            alan_bul(
                metin
            ),

        "aylik_kira":
            muhammen_bedel_bul(
                metin
            ),

        "gecici_teminat":
            gecici_teminat_bul(
                metin
            ),

        "ihale_tarihi":
            ihale_tarihi,

        "ihale_saati":
            ihale_saati,

        "veri_yontemi":
            "genel_belge",

        **yayin_araligi

    }


    return sonuc


def ihale_verisi_cikar(metin):

    tablo_verisi = tablo_satiri_bul(
        metin
    )


    if tablo_verisi:

        return tablo_verisi


    return genel_belge_verisi_cikar(
        metin
    )


def ana_program():

    print()
    print("==============================")
    print("KANTİN RADAR AI")
    print("MEB EK VERİ ÇIKARICI V6")
    print("==============================")


    if not os.path.exists(
        GIRDI_DOSYASI
    ):

        print(
            "HATA:",
            GIRDI_DOSYASI,
            "bulunamadı."
        )

        return False


    with open(
        GIRDI_DOSYASI,
        "r",
        encoding="utf-8"
    ) as dosya:

        sayfa_verisi = json.load(
            dosya
        )


    belge_metinleri = []

    okunan_dosyalar = []


    for ek in sayfa_verisi.get(
        "ek_dosyalar",
        []
    ):

        yerel_dosya = ek.get(
            "yerel_dosya",
            ""
        )


        if not yerel_dosya:

            continue


        if not os.path.exists(
            yerel_dosya
        ):

            continue


        uzanti = Path(
            yerel_dosya
        ).suffix.lower()


        if uzanti not in {
            ".pdf",
            ".docx"
        }:

            continue


        print()
        print(
            "Okunuyor:",
            yerel_dosya
        )


        metin = ek_dosya_oku(
            yerel_dosya
        )


        if not metin:

            print(
                "Metin çıkarılamadı."
            )

            continue


        metin_dosyasi = (
            metin_dosyasi_kaydet(
                yerel_dosya,
                metin
            )
        )


        print(
            "Okunan karakter:",
            len(
                metin
            )
        )


        print(
            "Metin dosyası:",
            metin_dosyasi
        )


        belge_metinleri.append(
            metin
        )


        okunan_dosyalar.append(
            {
                "dosya":
                    yerel_dosya,

                "karakter_sayisi":
                    len(
                        metin
                    ),

                "metin_dosyasi":
                    metin_dosyasi
            }
        )


    birlesik_metin = metin_temizle(
        "\n\n".join(
            belge_metinleri
        )
    )


    ihale = ihale_verisi_cikar(
        birlesik_metin
    )


    sonuc = {

        "kaynak":
            sayfa_verisi.get(
                "kaynak",
                "MEB"
            ),

        "sayfa_url":
            sayfa_verisi.get(
                "sayfa_url",
                ""
            ),

        "sayfa_basligi":
            sayfa_verisi.get(
                "sayfa_basligi",
                ""
            ),

        "yayin_tarihi":
            sayfa_verisi.get(
                "yayin_tarihi",
                ""
            ),

        "guncelleme_tarihi":
            sayfa_verisi.get(
                "guncelleme_tarihi",
                ""
            ),

        **ihale,

        "okunan_dosyalar":
            okunan_dosyalar,

        "ham_metin":
            birlesik_metin

    }


    with open(
        CIKTI_DOSYASI,
        "w",
        encoding="utf-8"
    ) as dosya:

        json.dump(
            sonuc,
            dosya,
            ensure_ascii=False,
            indent=4
        )


    print()
    print("==============================")
    print("ÇIKARILAN İHALE VERİSİ")
    print("==============================")


    print(
        "Okul:",
        sonuc.get(
            "okul_adi"
        )
        or "Bulunamadı"
    )


    print(
        "İlçe:",
        sonuc.get(
            "ilce"
        )
        or "Bulunamadı"
    )


    print(
        "İşyeri:",
        sonuc.get(
            "is_yeri_turu"
        )
        or "Bulunamadı"
    )


    print(
        "Öğrenci:",
        (
            sonuc.get(
                "ogrenci_sayisi"
            )
            if sonuc.get(
                "ogrenci_sayisi"
            )
            is not None
            else "Bulunamadı"
        )
    )


    print(
        "Personel:",
        (
            sonuc.get(
                "personel_sayisi"
            )
            if sonuc.get(
                "personel_sayisi"
            )
            is not None
            else "Bulunamadı"
        )
    )


    print(
        "Alan:",
        (
            sonuc.get(
                "alan_m2"
            )
            if sonuc.get(
                "alan_m2"
            )
            is not None
            else "Bulunamadı"
        )
    )


    print(
        "Aylık kira:",
        (
            sonuc.get(
                "aylik_kira"
            )
            if sonuc.get(
                "aylik_kira"
            )
            is not None
            else "Bulunamadı"
        )
    )


    print(
        "Geçici teminat:",
        (
            sonuc.get(
                "gecici_teminat"
            )
            if sonuc.get(
                "gecici_teminat"
            )
            is not None
            else "Bulunamadı"
        )
    )


    print(
        "İhale tarihi:",
        sonuc.get(
            "ihale_tarihi"
        )
        or "Bulunamadı"
    )


    print(
        "İhale saati:",
        sonuc.get(
            "ihale_saati"
        )
        or "Bulunamadı"
    )


    print(
        "İlan başlangıcı:",
        (
            sonuc.get(
                "ilan_baslangic_tarihi",
                ""
            )
            +
            " "
            +
            sonuc.get(
                "ilan_baslangic_saati",
                ""
            )
        ).strip()
        or "Bulunamadı"
    )


    print(
        "İlan bitişi:",
        (
            sonuc.get(
                "ilan_bitis_tarihi",
                ""
            )
            +
            " "
            +
            sonuc.get(
                "ilan_bitis_saati",
                ""
            )
        ).strip()
        or "Bulunamadı"
    )


    print(
        "Yöntem:",
        sonuc.get(
            "veri_yontemi",
            ""
        )
    )


    print()
    print(
        "JSON çıktısı:",
        CIKTI_DOSYASI
    )


    return True


if __name__ == "__main__":

    ana_program()